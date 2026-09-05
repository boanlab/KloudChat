"""Report export to `.docx`, `.pdf` and `.hwpx`, built from the artifact's own sections."""

from __future__ import annotations

import hashlib
import io
import logging
import re
import zipfile
from typing import Any

import PIL.Image
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus import Image as RLImage

from app.services import charts as chartkit
from app.services import design, doc_type, fonts, pictures, richtext

log = logging.getLogger(__name__)


def citation_text(source: dict[str, Any], style: str = "APA") -> str:
    """One stored source as a reference-list entry (APA, IEEE, MLA, Chicago); empty fields are skipped."""
    title = str(source.get("title") or "제목 없음").strip()
    author = str(source.get("author") or "").strip()
    publisher = str(source.get("publisher") or "").strip()
    year = str(source.get("year") or "").strip()
    url = str(source.get("url") or "").strip()
    style = style.upper()

    if style == "IEEE":
        ordinal = source.get("ordinal") or 1
        lead = f"[{ordinal}] " + (f"{author}, " if author else "")
        details = ", ".join(part for part in (publisher, year) if part)
        return (
            lead
            + f"“{title}.”"
            + (f" {details}." if details else "")
            + (f" [온라인]. {url}" if url else "")
        )
    if style == "MLA":
        parts = [f"{author}." if author else "", f"“{title}.”", publisher, year, url]
        return " ".join(part for part in parts if part).rstrip(".") + "."
    if style == "CHICAGO":
        lead = f"{author}. " if author else ""
        details = ", ".join(part for part in (publisher, year) if part)
        return (
            lead + f"“{title}.”" + (f" {details}." if details else "") + (f" {url}" if url else "")
        )

    # APA is the fallback for unknown styles.
    lead = f"{author}. " if author else ""
    dated = f"({year}). " if year else ""
    tail = ". ".join(part for part in (publisher, url) if part)
    return lead + dated + title + (f". {tail}" if tail else "")


def with_references(
    sections: list[dict], sources: list[dict] | None, style: str = "APA"
) -> list[dict]:
    """Return export sections with the report's reference list appended."""
    if not sources:
        return sections
    entries = [citation_text(source, style) for source in sources]
    return [
        *sections,
        {
            "id": "references",
            "heading": "참고문헌",
            "level": 1,
            "status": "done",
            "content": "\n".join(f"{index}. {entry}" for index, entry in enumerate(entries, 1)),
        },
    ]


#: A GFM table row.
_ROW = re.compile(r"^\s*\|(.+)\|\s*$")
#: A Markdown picture on its own line.
_IMAGE = re.compile(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$")
#: A fenced block, and the language on the fence.
_FENCE = re.compile(r"^\s*```\s*([A-Za-z0-9_-]*)\s*$")
_PAGE_BREAK = re.compile(r"^\s*<!--\s*pagebreak\s*-->\s*$", re.I)


#: A footnote mark inside a line of prose, in the same notation.
_MARK = re.compile(r"\[\^([^\]]{1,16})\]")

#: `1` → `¹`: plain characters need no OWPML run markup.
_SUPERSCRIPT = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")


def _raised_marks(text: str) -> str:
    """`[^1]` → `¹`, for the formats with no footnote model."""

    def raise_one(found: re.Match[str]) -> str:
        mark = found.group(1)
        return mark.translate(_SUPERSCRIPT) if mark.isdigit() else f"({mark})"

    return _MARK.sub(raise_one, text)


def _marks_in(text: str) -> set[str]:
    """The footnote numbers a line cites."""
    return {found.group(1) for found in _MARK.finditer(text)}


#: A footnote body on its own line, in GFM notation.
_NOTE = re.compile(r"^\[\^([^\]]{1,16})\]:\s*(.+)$")


def _chart_as_rows(chart: dict) -> list[list[str]]:
    """A chart's numbers as a table, for the formats that cannot draw one."""
    head = ["항목", *[name or "값" for name, _ in chart["series"]]]
    body = [
        [category, *[_number(series[row]) for _, series in chart["series"]]]
        for row, category in enumerate(chart["categories"])
    ]
    return [head, *body]


def _number(value: float) -> str:
    """A value without the decimal point it did not come with."""
    return f"{value:,.0f}" if float(value).is_integer() else f"{value:,.2f}".rstrip("0")


def _chart_block(source: str) -> dict | None:
    """A `chart` fence as its numbers, or `None`.

    Line 1: `kind | unit`. Line 2: `axis | series…`. Then one row per
    category with one value per series; short rows are dropped, not padded.
    """
    lines = [line.strip() for line in (source or "").splitlines() if line.strip()]
    if len(lines) < 3:
        return None

    head = [cell.strip() for cell in lines[0].split("|")]
    kind = head[0].lower() if head else "bar"
    if kind not in ("bar", "line"):
        kind = "bar"
    unit = head[1] if len(head) > 1 else ""

    names = [cell.strip() for cell in lines[1].split("|")][1:]
    if not names:
        return None

    categories: list[str] = []
    columns: list[list[float]] = [[] for _ in names]
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split("|")]
        if len(cells) < len(names) + 1 or not cells[0]:
            continue
        try:
            values = [float(cell.replace(",", "")) for cell in cells[1 : len(names) + 1]]
        except ValueError:
            continue
        categories.append(cells[0])
        for column, value in zip(columns, values, strict=True):
            column.append(value)

    # Two points make a shape; one is a number, and there is a block for that.
    if len(categories) < 2:
        return None
    return {
        "kind": kind,
        "unit": unit,
        "categories": categories[:8],
        "series": [(name, column[:8]) for name, column in zip(names[:2], columns, strict=False)],
    }


def _kpi_rows(source: str, *, limit: int = 4) -> list[tuple[str, str]]:
    """`[(값, 이름)]` from a `kpi` or `steps` fence: one pair per line, split on the first `|`.

    Lines without a separator are dropped. `limit` is 4 for a kpi strip, 8 for steps.
    """
    out: list[tuple[str, str]] = []
    for line in (source or "").splitlines():
        if "|" not in line:
            continue
        value, _, label = line.partition("|")
        value, label = value.strip(), label.strip()
        if value and label:
            out.append((value, label))
    return out[:limit]


def _in_pairs(
    cards: list[tuple[str, list[str]]],
) -> list[list[tuple[str, list[str]]]]:
    """The cards two at a time; an odd last card gets an empty partner."""
    padded: list[tuple[str, list[str]]] = list(cards)
    if len(padded) % 2:
        padded.append(("", []))
    return [padded[i : i + 2] for i in range(0, len(padded), 2)]


def _cards(source: str, *, limit: int = 6) -> list[tuple[str, list[str]]]:
    """`[(제목, [줄, …])]` from a `cards` fence: `## 제목` then `- 줄` lines; six cards at most."""
    cards: list[tuple[str, list[str]]] = []
    for raw in (source or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#"):
            cards.append((line.lstrip("#").strip(), []))
        elif cards:
            cards[-1][1].append(line.lstrip("-*").strip() or line)
    return [(title, items) for title, items in cards if title][:limit]


def _callout(source: str) -> tuple[str, list[str]]:
    """`(제목, [줄, …])` from a `callout` fence; the first line is the title."""
    lines = [raw.strip() for raw in (source or "").splitlines() if raw.strip()]
    if not lines:
        return ("", [])
    return (lines[0].lstrip("#").strip(), [line.lstrip("-*").strip() for line in lines[1:]])


def diagram_key(source: str) -> str:
    """A stable key for a diagram's source: the browser stores its raster under it, exporters look it up.

    Whitespace-insensitive, since the browser and the server normalise differently.
    """
    normalised = "\n".join(line.rstrip() for line in (source or "").strip().splitlines())
    return hashlib.sha256(normalised.encode()).hexdigest()[:16]


_RULE = re.compile(r"^\s*\|[\s:|-]+\|\s*$")


def _cells(line: str) -> list[str]:
    """One row's cells. An escaped pipe stays inside its cell."""
    body = _ROW.match(line).group(1)  # type: ignore[union-attr]
    return [cell.replace("\\|", "|").strip() for cell in re.split(r"(?<!\\)\|", body)]


def _as_grid(rows: richtext.Grid | list[list[str]]) -> richtext.Grid:
    """Whatever a caller has, as a grid. Plain rows merge nothing."""
    if isinstance(rows, richtext.Grid):
        return rows
    return richtext.Grid(
        rows=[[richtext.Cell(cell.replace("<br>", "\n")) for cell in row] for row in rows]
    )


def _matched(
    written: list[list[str]], grids: list[richtext.Grid], used: set[int]
) -> richtext.Grid | None:
    """The grid a GFM table was written from, matched on text rather than position."""
    for index, grid in enumerate(grids):
        if index in used:
            continue
        if grid.flat(newline="<br>") == written:
            used.add(index)
            return grid
    return None


#: The marker each list depth hangs: `•`, then the sub-markers a Korean 공문 uses.
_MARKERS = ("•", "–", "·")
#: How far one level is set in. Points here; each writer converts.
_INDENT = 14.0


def _markdown_to_lines(
    text: str, grids: list[richtext.Grid] | None = None
) -> list[tuple[str, Any, str, int]]:
    """`(kind, payload, marker, depth)` per line.

    Kinds: heading, bullet, number, body, note, table, image, chart, steps,
    kpi, cards, callout, diagram, pagebreak. A `table` carries its grid as the
    payload. `marker` is `•` or `3.`; `depth` is the indent level (two spaces
    each). Numbering follows Markdown: the first item's number starts the run.
    """
    out: list[tuple[str, Any, str, int]] = []
    number = 0
    #: Numbering per list depth; deeper levels are cleared when a level closes.
    counts: dict[int, int] = {}
    rows: list[list[str]] = []
    #: Grids for this section's tables, spent as they are matched.
    spare = list(grids or [])
    used: set[int] = set()
    #: Whether the table being collected has had its rule row yet.
    ruled = False
    #: The fenced block being collected, and what language it claimed.
    fence: list[str] | None = None
    fence_lang = ""

    def close_table() -> None:
        nonlocal rows, ruled
        ruled = False
        if rows:
            width = max(len(row) for row in rows)
            padded = [row + [""] * (width - len(row)) for row in rows]
            found = _matched(padded, spare, used)
            out.append(("table", found or _as_grid(padded), "", 0))
            rows = []

    for raw in (text or "").splitlines():
        line = raw.rstrip()
        if fence is not None:
            if _FENCE.match(line):
                source = "\n".join(fence)
                if fence_lang.lower() == "chart":
                    if drawn := _chart_block(source):
                        # The source rides in the marker slot; `.hwpx` looks up
                        # the browser's raster by its digest.
                        out.append(("chart", drawn, source, 0))
                elif fence_lang.lower() == "steps":
                    if steps := _kpi_rows(source, limit=8):
                        out.append(("steps", steps, "", 0))
                elif fence_lang.lower() == "kpi":
                    if figures := _kpi_rows(source):
                        out.append(("kpi", figures, "", 0))
                elif fence_lang.lower() == "cards":
                    if cards := _cards(source):
                        out.append(("cards", cards, "", 0))
                elif fence_lang.lower() == "callout":
                    if callout := _callout(source):
                        if callout[0]:
                            out.append(("callout", callout, "", 0))
                elif fence_lang.lower() == "mermaid":
                    # Placed by key; the exporters look up the browser's raster.
                    out.append(("diagram", {"source": source, "key": diagram_key(source)}, "", 0))
                else:
                    # Any other fence goes out as prose.
                    out.extend(("body", one, "", 0) for one in fence if one.strip())
                fence = None
                fence_lang = ""
            else:
                fence.append(raw)
            continue
        if found := _FENCE.match(line):
            close_table()
            fence = []
            fence_lang = found.group(1)
            continue
        if _PAGE_BREAK.match(line):
            close_table()
            number = 0
            counts.clear()
            out.append(("pagebreak", "", "", 0))
            continue
        if not line.strip():
            # A blank line does not end a table: hand-edited documents carry gaps.
            if not rows:
                close_table()
            continue
        if _RULE.match(line) and rows:
            # A second rule row starts a new table whose head is the row just read.
            if ruled:
                head = rows.pop()
                close_table()
                rows.append(head)
            ruled = True
            continue
        if _ROW.match(line):
            number = 0
            counts.clear()
            rows.append(_cells(line))
            continue
        close_table()
        if picture := _IMAGE.match(line):
            # Only embedded pictures are carried: an exporter must not make a network call.
            decoded = pictures.decode(picture.group(2).strip())
            if decoded:
                mime, data = decoded
                out.append(
                    ("image", {"data": data, "mime": mime, "caption": picture.group(1)}, "", 0)
                )
            continue
        if heading := re.match(r"^#{2,6}\s+(.*)$", line):
            number = 0
            counts.clear()
            out.append(("heading", heading.group(1).strip(), "", 0))
        elif note := _NOTE.match(line):
            number = 0
            counts.clear()
            out.append(("note", note.group(2).strip(), note.group(1), 0))
        elif bullet := re.match(r"^(\s*)[-*+]\s+(.*)$", line):
            number = 0
            counts.clear()
            depth = min(len(bullet.group(1)) // 2, len(_MARKERS) - 1)
            out.append(("bullet", bullet.group(2).strip(), _MARKERS[depth], depth))
        elif numbered := re.match(r"^(\s*)(\d{1,9})[.)]\s+(.*)$", line):
            depth = min(len(numbered.group(1)) // 2, len(_MARKERS) - 1)
            # Counted per level so a sub-list does not advance its parent.
            # `is None`, not falsy: a list may legally start at 0.
            counts[depth] = (
                counts[depth] + 1 if counts.get(depth) is not None else int(numbered.group(2))
            )
            for deeper in [d for d in counts if d > depth]:
                del counts[deeper]
            number = counts[depth]
            out.append(("number", numbered.group(3).strip(), f"{number}.", depth))
        else:
            number = 0
            counts.clear()
            out.append(("body", line.strip(), "", 0))
    close_table()
    if fence:
        # An unclosed fence is a truncated document, not a diagram.
        out.extend(("body", one, "", 0) for one in fence if one.strip())
    return out


def _strip_inline(text: str) -> str:
    """Bold, italic and code markers removed rather than rendered.

    A code span must not open or close on a space, so a lone backtick — the
    Korean date `28.03 — survives.
    """
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\w)\*(.+?)\*(?!\w)", r"\1", text)
    return re.sub(r"`(\S(?:[^`]*\S)?)`", r"\1", text)


def _format_key(text: str) -> str:
    return re.sub(r"\s+", " ", _strip_inline(str(text))).strip()


def _next_format(blocks: list[dict], text: str, cursor: int) -> tuple[dict | None, int]:
    """Match one prose line to the next editable-HTML block, in document order."""
    wanted = _format_key(text)
    for index in range(cursor, len(blocks)):
        block = blocks[index]
        if _format_key(str(block.get("text") or "")) == wanted:
            return block, index + 1
    return None, cursor


def _docx_styled(paragraph, block: dict, fallback: str) -> None:
    """Write one editable HTML block as real Word paragraph/run properties."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    block_style = block.get("style") or {}
    alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(str(block_style.get("text-align") or "").lower())
    if alignment is not None:
        paragraph.alignment = alignment
    try:
        if line_height := block_style.get("line-height"):
            paragraph.paragraph_format.line_spacing = float(line_height)
    except (TypeError, ValueError):
        pass

    runs = block.get("runs") or [{"text": fallback, "style": {}}]
    for item in runs:
        run = paragraph.add_run(str(item.get("text") or ""))
        style = item.get("style") or {}
        weight = str(style.get("font-weight") or "").lower()
        run.bold = weight in {"bold", "600", "700", "800", "900"}
        run.italic = str(style.get("font-style") or "").lower() == "italic"
        decoration = str(style.get("text-decoration") or "").lower()
        run.underline = "underline" in decoration
        run.font.strike = "line-through" in decoration
        if family := str(style.get("font-family") or "").split(",")[0].strip(" '\""):
            run.font.name = family
        size = str(style.get("font-size") or "")
        if found := re.fullmatch(r"(\d+(?:\.\d+)?)pt", size, re.I):
            run.font.size = Pt(max(6, min(72, float(found.group(1)))))
        colour = str(style.get("color") or "")
        if found := re.fullmatch(r"#([0-9a-f]{6})", colour, re.I):
            run.font.color.rgb = RGBColor.from_string(found.group(1).upper())
        elif found := re.fullmatch(r"#([0-9a-f]{3})", colour, re.I):
            value = "".join(character * 2 for character in found.group(1))
            run.font.color.rgb = RGBColor.from_string(value.upper())
        background = str(style.get("background-color") or "")
        if found := re.fullmatch(r"#([0-9a-f]{6}|[0-9a-f]{3})", background, re.I):
            value = found.group(1)
            if len(value) == 3:
                value = "".join(character * 2 for character in value)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), value.upper())
            run._r.get_or_add_rPr().append(shading)


def _pdf_styled(block: dict, fallback: str) -> str:
    """ReportLab-safe markup for the formatting the document editor stores."""
    out: list[str] = []
    for item in block.get("runs") or [{"text": fallback, "style": {}}]:
        text = _escape(str(item.get("text") or "")).replace("\n", "<br/>")
        style = item.get("style") or {}
        tags: list[tuple[str, str]] = []
        attributes: list[str] = []
        size = str(style.get("font-size") or "")
        if found := re.fullmatch(r"(\d+(?:\.\d+)?)pt", size, re.I):
            attributes.append(f'size="{max(6, min(72, float(found.group(1)))):g}"')
        for source, target in (("color", "color"), ("background-color", "backColor")):
            value = str(style.get(source) or "")
            if re.fullmatch(r"#[0-9a-f]{3}(?:[0-9a-f]{3})?", value, re.I):
                attributes.append(f'{target}="{value}"')
        if attributes:
            tags.append(("font", " ".join(attributes)))
        weight = str(style.get("font-weight") or "").lower()
        if weight in {"bold", "600", "700", "800", "900"}:
            tags.append(("b", ""))
        if str(style.get("font-style") or "").lower() == "italic":
            tags.append(("i", ""))
        decoration = str(style.get("text-decoration") or "").lower()
        if "underline" in decoration:
            tags.append(("u", ""))
        if "line-through" in decoration:
            tags.append(("strike", ""))
        for tag, attributes in tags:
            text = f"<{tag}{(' ' + attributes) if attributes else ''}>{text}</{tag}>"
        out.append(text)
    return "".join(out)


def _pdf_block_style(base: ParagraphStyle, block: dict) -> ParagraphStyle:
    style = block.get("style") or {}
    alignment = {
        "left": TA_LEFT,
        "center": TA_CENTER,
        "right": TA_RIGHT,
        "justify": TA_JUSTIFY,
    }.get(str(style.get("text-align") or "").lower(), base.alignment)
    leading = base.leading
    try:
        if line_height := style.get("line-height"):
            leading = base.fontSize * float(line_height)
    except (TypeError, ValueError):
        pass
    return ParagraphStyle(
        f"{base.name}-edited-{id(block)}", parent=base, alignment=alignment, leading=leading
    )


def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    """A Word field as `w:fldSimple`; python-docx has no API for one."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    if placeholder:
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = placeholder
        run.append(text)
        field.append(run)
    paragraph._p.append(field)


def _page_setup(document, settings: dict | None = None) -> None:
    """A4 with 20mm margins by default; python-docx starts on US Letter."""
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm

    settings = settings or {}
    margins = settings.get("margins") if isinstance(settings.get("margins"), dict) else {}
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(float(margins.get("left", 20)))
        section.right_margin = Mm(float(margins.get("right", 20)))
        section.top_margin = Mm(float(margins.get("top", 20)))
        section.bottom_margin = Mm(float(margins.get("bottom", 20)))
        header_text = str(settings.get("header") or "")
        if header_text:
            header = section.header.paragraphs[0]
            header.text = header_text
        if not bool(settings.get("firstPageHeader", False)):
            section.different_first_page_header_footer = True
            section.first_page_header.paragraphs[0].text = ""
        # Page number centred in the footer, as a Korean submission expects.
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = str(settings.get("footer") or "")
        if footer_text:
            footer.add_run(footer_text + " · ")
        page_numbers = str(settings.get("pageNumbers") or "page-total")
        if page_numbers != "none":
            _field(footer, "PAGE", "1")
            if page_numbers == "page-total":
                footer.add_run(" / ")
                _field(footer, "NUMPAGES", "1")


#: A table of contents only for a document long enough to need one.
_TOC_SECTIONS = 6
_TOC_CHARS = 6_000


def _wants_toc(sections: list[dict]) -> bool:
    length = sum(len(str(s.get("content") or "")) for s in sections)
    return len(sections) >= _TOC_SECTIONS and length >= _TOC_CHARS


def _update_fields_on_open(document) -> None:
    """Asks Word to refresh fields on open, so the 목차 fills itself in."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings_element = document.settings.element
    if settings_element.find(qn("w:updateFields")) is None:
        flag = OxmlElement("w:updateFields")
        flag.set(qn("w:val"), "true")
        settings_element.append(flag)


def _docx_type_scale(document) -> None:
    """Word's built-in styles at the document scale, so the file matches the page view.

    Only on a plain document: a 서식's own Word file keeps its styles.
    """
    from docx.shared import Pt

    sizes = {
        "Title": ("title", True),
        "Heading 1": ("h1", True),
        "Heading 2": ("h2", True),
        "Heading 3": ("h3", True),
        "Normal": ("body", False),
        "Body Text": ("body", False),
        "List Bullet": ("body", False),
        "List Number": ("body", False),
        "Caption": ("caption", False),
    }
    for name, (key, bold) in sizes.items():
        try:
            style = document.styles[name]
        except KeyError:
            continue
        style.font.size = Pt(doc_type.TYPE[key])
        if bold:
            style.font.bold = True
        if key == "body":
            style.paragraph_format.line_spacing = doc_type.LEADING["body"]


def _korean_fonts(document) -> None:
    """맑은 고딕 for Hangul and Calibri for Latin on the default styles; python-docx names no East Asian font."""
    from docx.oxml.ns import qn

    for name in ("Normal", "Body Text", "Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = document.styles[name]
        except KeyError:
            continue
        style.font.name = "Calibri"
        # Headings inherit a blue theme colour from Word's defaults; the accent or ink
        # is set per heading, so the style itself goes neutral.
        if name.startswith("Heading") or name == "Title":
            style.font.color.rgb = None
        rpr = style.element.get_or_add_rPr()
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            from docx.oxml import OxmlElement

            fonts = OxmlElement("w:rFonts")
            rpr.append(fonts)
        fonts.set(qn("w:eastAsia"), "맑은 고딕")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")


def _table_of_contents(document) -> None:
    """A 목차 field over the headings; shows its placeholder until Word updates it."""
    from docx.shared import Pt

    heading = document.add_paragraph("목차", style="Heading 1")
    heading.paragraph_format.space_after = Pt(6)
    _field(
        document.add_paragraph(),
        'TOC \\o "1-2" \\h \\z \\u',
        "목차를 보려면 이 줄에서 F9 를 누르세요.",
    )
    # No page break after it: an unexpanded field would leave the first page empty.


def _docx_cell_margins(table, vertical_pt: float = 3.0, horizontal_pt: float = 6.0) -> None:
    """Room between a cell's border and its text; Word's default is a hair."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = table._tbl.tblPr
    margins = properties.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        properties.append(margins)
    for side, points in (
        ("top", vertical_pt),
        ("bottom", vertical_pt),
        ("left", horizontal_pt),
        ("right", horizontal_pt),
    ):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(int(points * 20)))
        node.set(qn("w:type"), "dxa")


def _docx_table(document, rows: richtext.Grid | list[list[str]], accent: str = "") -> None:
    """A real Word table, head row emphasised, merges merged.

    Built as a full rectangle and merged afterwards; text is written after the
    merge, since a merged cell keeps the text of every cell that went into it.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    grid = richtext.Grid(
        rows=[row for row in _as_grid(rows).rows if any(c.text.strip() for c in row)]
    )
    if not grid.rows:
        return
    width = grid.width
    table = document.add_table(rows=len(grid.rows), cols=width)
    table.style = "Table Grid"
    _docx_cell_margins(table)

    covered: set[tuple[int, int]] = set()
    for r, row in enumerate(grid.rows):
        column = 0
        for source in row:
            while (r, column) in covered:
                column += 1
            if column >= width:
                break
            across = min(source.colspan, width - column)
            down = min(source.rowspan, len(grid.rows) - r)
            for dr in range(down):
                for dc in range(across):
                    if (dr, dc) != (0, 0):
                        covered.add((r + dr, column + dc))
            cell = table.cell(r, column)
            if across > 1 or down > 1:
                cell = cell.merge(table.cell(r + down - 1, column + across - 1))
            lines = [_strip_inline(line) for line in source.text.split("\n")] or [""]
            cell.text = lines[0]
            for extra in lines[1:]:
                cell.add_paragraph(extra)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(doc_type.TYPE["table"])
            if r == 0:
                if accent:
                    properties = cell._tc.get_or_add_tcPr()
                    shade = OxmlElement("w:shd")
                    shade.set(qn("w:fill"), "F2F2F2")
                    properties.append(shade)
                    borders = properties.find(qn("w:tcBorders"))
                    if borders is None:
                        borders = OxmlElement("w:tcBorders")
                        properties.append(borders)
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "10")
                    bottom.set(qn("w:color"), accent.lstrip("#").upper())
                    borders.append(bottom)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            column += across
    document.add_paragraph().paragraph_format.space_after = Pt(4)


#: Separator notes (ids -1 and 0) every footnotes part must open with, or Word repairs the file.
_DOCX_FOOTNOTE_HEAD = (
    '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0"'
    ' w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing'
    ' w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r>'
    "<w:continuationSeparator/></w:r></w:p></w:footnote>"
)


class _Footnotes:
    """The document's footnote part, written by hand: python-docx has no footnote API.

    A `footnotes.xml` Word does not accept produces a repair prompt on open.
    """

    def __init__(self) -> None:
        self.notes: list[str] = []

    def write(self, paragraph, text: str, notes: dict[str, str]) -> None:
        """A line of prose with its marks placed as real references; a mark with no note is dropped."""
        cursor = 0
        for found in _MARK.finditer(text):
            if before := text[cursor : found.start()]:
                paragraph.add_run(before)
            if note := notes.get(found.group(1)):
                self.add(paragraph, note)
            cursor = found.end()
        if rest := text[cursor:]:
            paragraph.add_run(rest)

    def add(self, paragraph, text: str) -> None:
        """One note, and the mark in the prose that points at it."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        self.notes.append(text)
        run = paragraph.add_run()
        # `FootnoteReference` is the character style every template carries.
        properties = OxmlElement("w:rPr")
        style = OxmlElement("w:rStyle")
        style.set(qn("w:val"), "FootnoteReference")
        properties.append(style)
        vertical = OxmlElement("w:vertAlign")
        vertical.set(qn("w:val"), "superscript")
        properties.append(vertical)
        run._r.append(properties)
        reference = OxmlElement("w:footnoteReference")
        # Ids -1 and 0 are the separators, so the first real note is 1.
        reference.set(qn("w:id"), str(len(self.notes)))
        run._r.append(reference)

    def attach(self, document) -> None:
        """Writes the part and wires it to the document, if there are any."""
        if not self.notes:
            return
        from docx.opc.constants import RELATIONSHIP_TYPE as RT  # noqa: N814
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part

        body = "".join(
            f'<w:footnote w:id="{index}"><w:p><w:pPr>'
            '<w:pStyle w:val="FootnoteText"/></w:pPr><w:r><w:rPr>'
            '<w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r>'
            f'<w:r><w:t xml:space="preserve"> {_escape(note)}</w:t></w:r>'
            "</w:p></w:footnote>"
            for index, note in enumerate(self.notes, start=1)
        )
        xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            + _DOCX_FOOTNOTE_HEAD
            + body
            + "</w:footnotes>"
        )
        part = Part(
            PackURI("/word/footnotes.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
            xml.encode("utf-8"),
            document.part.package,
        )
        document.part.relate_to(part, RT.FOOTNOTES)


#: The inline drawing frame that puts a chart part into a paragraph.
_DOCX_CHART_DRAWING = (
    '<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
    ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
    ' xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart"'
    ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
    '<wp:inline distT="0" distB="0" distL="0" distR="0">'
    '<wp:extent cx="{cx}" cy="{cy}"/>'
    '<wp:docPr id="{n}" name="Chart {n}"/>'
    "<a:graphic><a:graphicData"
    ' uri="http://schemas.openxmlformats.org/drawingml/2006/chart">'
    '<c:chart r:id="{rid}"/>'
    "</a:graphicData></a:graphic>"
    "</wp:inline></w:drawing>"
)


def _docx_chart(document, chart: dict, index: int, style: dict | None) -> bool:
    """A native Word chart, or `False` when it could not be built.

    The chart part, its document relationship, its workbook relationship and
    both content types must all agree, or Word offers to repair the file.
    """
    from docx.opc.constants import CONTENT_TYPE as CT
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part
    from docx.oxml import parse_xml
    from pptx.dml.color import RGBColor

    built = chartkit.part(
        chart["kind"],
        chart["categories"],
        chart["series"],
        unit=chart.get("unit") or "",
        accent=RGBColor.from_string((style["accent"] if style else "#5b5bd6").lstrip("#").upper()),
        muted=RGBColor.from_string((style["muted"] if style else "#666666").lstrip("#").upper()),
        faces=chartkit.FACES["serif" if style and style.get("font") == "serif" else "gothic"],
    )
    if built is None:
        return False
    chart_xml, workbook = built

    try:
        package = document.part.package
        chart_part = Part(
            PackURI(f"/word/charts/chart{index}.xml"), CT.DML_CHART, chart_xml, package
        )
        # Without the workbook relationship the chart has no editable data.
        chart_part.relate_to(
            Part(
                PackURI(f"/word/embeddings/chart{index}.xlsx"),
                CT.SML_SHEET,
                workbook,
                package,
            ),
            RT.PACKAGE,
        )
        relationship_id = document.part.relate_to(chart_part, RT.CHART)
    except Exception as exc:  # noqa: BLE001 — a chart is not worth a failed export
        log.warning("could not attach a chart to the docx: %s", exc)
        return False

    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # centred, like a figure
    run = paragraph.add_run()
    run._r.append(
        parse_xml(
            _DOCX_CHART_DRAWING.format(
                # 150mm × 85mm in EMU.
                cx=int(150 * 36000),
                cy=int(85 * 36000),
                n=index,
                rid=relationship_id,
            )
        )
    )
    return True


def _docx_steps(document, steps: list[tuple[str, str]]) -> None:
    """A numbered procedure as a two-column table; numbers are literal text so they never renumber."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    if not steps:
        return
    table = document.add_table(rows=len(steps), cols=2)
    table.style = "Table Grid"
    for row, (name, detail) in enumerate(steps):
        left = table.cell(row, 0).paragraphs[0]
        left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        number = left.add_run(str(row + 1))
        number.bold = True
        number.font.size = Pt(12)

        right = table.cell(row, 1).paragraphs[0]
        title = right.add_run(name)
        title.bold = True
        if detail:
            right.add_run(f"  {detail}")
        table.rows[row].cells[0].width = Mm(12)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def _docx_kpi(document, figures: list[tuple[str, str]]) -> None:
    """A row of figures as a borderless Word table: values large, labels small."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    if not figures:
        return
    table = document.add_table(rows=2, cols=len(figures))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # No `Table Grid`: a strip is separated by space, not lines.
    for column, (value, label) in enumerate(figures):
        top = table.cell(0, column).paragraphs[0]
        top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = top.add_run(value)
        run.bold = True
        run.font.size = Pt(doc_type.TYPE["kpi"])

        bottom = table.cell(1, column).paragraphs[0]
        bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note = bottom.add_run(label)
        note.font.size = Pt(doc_type.TYPE["kpiLabel"])
        note.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def _docx_cards(document, cards: list[tuple[str, list[str]]]) -> None:
    """A card grid as a two-column Word table."""
    from docx.shared import Pt

    if not cards:
        return
    rows = (len(cards) + 1) // 2
    table = document.add_table(rows=rows, cols=2)
    table.style = "Table Grid"
    for index, (title, items) in enumerate(cards):
        cell = table.cell(index // 2, index % 2)
        head = cell.paragraphs[0]
        head.paragraph_format.space_after = Pt(3)
        run = head.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        for item in items:
            line = cell.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            written = line.add_run(f"· {item}")
            written.font.size = Pt(9.5)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def _docx_callout(document, callout: tuple[str, list[str]]) -> None:
    """The boxed line, as a one-cell table so the border comes with it."""
    from docx.shared import Pt

    title, lines = callout
    if not title:
        return
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    head = cell.paragraphs[0]
    head.paragraph_format.space_after = Pt(2)
    run = head.add_run(title)
    run.bold = True
    for line in lines:
        written = cell.add_paragraph()
        written.paragraph_format.space_after = Pt(2)
        written.add_run(line).font.size = Pt(9.5)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def _diagram_picture(section: dict, drawn: dict) -> dict | None:
    """The raster a browser stored for this diagram, or `None` when nobody has opened the document yet."""
    store = section.get("diagrams")
    if not isinstance(store, dict):
        return None
    decoded = pictures.decode(str(store.get(drawn.get("key")) or ""))
    if not decoded:
        return None
    mime, data = decoded
    return {"data": data, "mime": mime, "caption": ""}


def _docx_picture(document, picture: dict) -> None:
    """One figure, centred, with its caption under it."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    data = picture.get("data")
    if not data:
        return
    width_pt, height_pt = _picture_size(data)
    try:
        # Both dimensions: from the width alone Word lets a portrait picture fill the page.
        document.add_picture(io.BytesIO(data), width=Pt(width_pt), height=Pt(height_pt))
    except Exception as exc:  # noqa: BLE001 — a bad picture is not a failed export
        log.warning("could not place a picture in the docx: %s", exc)
        return
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = str(picture.get("caption") or "")
    if caption:
        paragraph = document.add_paragraph(caption)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(8)
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


#: Body paragraphs go in `Body Text`, not `Normal`, so a 서식's body settings
#: apply without touching headings.
_BODY = "Body Text"


def to_docx(
    title: str,
    sections: list[dict],
    *,
    tokens: dict[str, str] | None = None,
    template: str = "",
    page_settings: dict | None = None,
) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    style = design.normalise_tokens(tokens) if tokens else None
    #: Colour only: a run-level typeface would override the reviewer's template.
    accent = RGBColor.from_string(style["accent"].lstrip("#").upper()) if style else None

    def recolour(heading) -> None:
        if accent is None:
            return
        for run in heading.runs:
            run.font.color.rgb = accent

    # The 서식's own Word file makes the styles, page and theme the template's.
    # An unreadable template falls back to defaults rather than failing.
    document = None
    if template:
        try:
            document = Document(template)
        except Exception as exc:  # noqa: BLE001 — a plain export beats none
            log.warning("docx template unreadable (%s): %s", template, exc)
    document_is_plain = document is None
    if document is None:
        document = Document()
        _page_setup(document, page_settings)
    elif page_settings:
        _page_setup(document, page_settings)
    visual_style = (style or {}).get("visualStyle") or "editorial"

    def visualise_heading(paragraph, level: int) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        if visual_style == "editorial":
            recolour(paragraph)
            if level == 0 and accent is not None:
                properties = paragraph._p.get_or_add_pPr()
                borders = properties.find(qn("w:pBdr"))
                if borders is None:
                    borders = OxmlElement("w:pBdr")
                    properties.append(borders)
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "10")
                bottom.set(qn("w:space"), "6")
                bottom.set(qn("w:color"), (style or {})["accent"].lstrip("#").upper())
                borders.append(bottom)
            return

        for run in paragraph.runs:
            if visual_style == "poster":
                run.font.size = Pt(doc_type.TYPE["title" if level == 0 else "h1"])
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if level == 0 else accent
                run.font.bold = True
            else:
                run.font.size = Pt(doc_type.TYPE["title" if level == 0 else "h1"])
                run.font.color.rgb = (
                    RGBColor(0x1A, 0x1A, 0x1A) if level == 0 else RGBColor(0x66, 0x66, 0x66)
                )
                run.font.bold = level == 0
        if visual_style == "poster" and level == 0:
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), (style or {}).get("accent", "#5b5bd6").lstrip("#"))
            paragraph._p.get_or_add_pPr().append(shade)

    title_heading = document.add_heading(title, level=0)
    visualise_heading(title_heading, 0)
    if _wants_toc(sections):
        _table_of_contents(document)
        _update_fields_on_open(document)
    if document_is_plain:
        _korean_fonts(document)
        _docx_type_scale(document)

    footnotes = _Footnotes()
    #: Charts are numbered across the document; each is its own part.
    charts = 0
    for section in sections:
        section_heading = document.add_heading(section.get("heading") or "", level=1)
        visualise_heading(section_heading, 1)
        lines = _markdown_to_lines(section.get("content") or "", section.get("tables"))
        formatted = list(section.get("_formatting") or [])
        format_cursor = 0
        #: Notes by mark, read ahead: a note is written after the sentence that cites it.
        notes = {mark: str(body) for kind, body, mark, _ in lines if kind == "note"}
        cited: set[str] = set()
        for kind, text, marker, depth in lines:
            if kind == "pagebreak":
                document.add_page_break()
                continue
            if kind == "table":
                _docx_table(document, text, (style or {}).get("accent", ""))
                continue
            if kind == "image":
                _docx_picture(document, text)
                continue
            if kind == "kpi":
                _docx_kpi(document, text)
                continue
            if kind == "steps":
                _docx_steps(document, text)
                continue
            if kind == "cards":
                _docx_cards(document, text)
                continue
            if kind == "callout":
                _docx_callout(document, text)
                continue
            if kind == "chart":
                charts += 1
                if not _docx_chart(document, text, charts, style):
                    # Fallback: the numbers as a table.
                    _docx_table(document, _chart_as_rows(text), (style or {}).get("accent", ""))
                continue
            if kind == "diagram":
                if picture := _diagram_picture(section, text):
                    _docx_picture(document, picture)
                else:
                    note = document.add_paragraph(
                        "[다이어그램 — 화면에서 한 번 열면 그림으로 저장됩니다]", style=_BODY
                    )
                    note.paragraph_format.space_after = Pt(6)
                continue
            clean = _strip_inline(text)
            if kind == "note":
                # Placed at its mark already; an uncited note is written as prose.
                if marker not in cited:
                    document.add_paragraph(clean, style=_BODY)
                continue
            block, format_cursor = _next_format(formatted, clean, format_cursor)
            if kind == "heading":
                paragraph = document.add_heading("", level=2)
                if block:
                    _docx_styled(paragraph, block, clean)
                else:
                    paragraph.add_run(clean)
            elif kind == "bullet":
                paragraph = document.add_paragraph(style="List Bullet")
                if block and not _MARK.search(clean):
                    _docx_styled(paragraph, block, clean)
                else:
                    footnotes.write(paragraph, clean, notes)
                if depth:
                    paragraph.paragraph_format.left_indent = Inches(0.25 * (depth + 1))
                cited |= _marks_in(clean)
            elif kind == "number":
                # Not "List Number": its numbering runs on across separate lists.
                paragraph = document.add_paragraph(f"{marker} ", style=_BODY)
                if block and not _MARK.search(clean):
                    _docx_styled(paragraph, block, clean)
                else:
                    footnotes.write(paragraph, clean, notes)
                paragraph.paragraph_format.left_indent = Inches(0.25 * (depth + 1))
                paragraph.paragraph_format.space_after = Pt(3)
                cited |= _marks_in(clean)
            else:
                paragraph = document.add_paragraph(style=_BODY)
                if block and not _MARK.search(clean):
                    _docx_styled(paragraph, block, clean)
                else:
                    footnotes.write(paragraph, clean, notes)
                paragraph.paragraph_format.space_after = Pt(6)
                cited |= _marks_in(clean)
        # Figures attached to the section itself.
        for picture in section.get("images") or []:
            _docx_picture(document, picture)

    footnotes.attach(document)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _at_depth(style, depth: int):
    """A list style set in by `depth` levels. Cloned: reportlab styles are shared objects."""
    if not depth:
        return style
    return ParagraphStyle(
        f"{style.name}-{depth}",
        parent=style,
        leftIndent=style.leftIndent + depth * (_INDENT / 72 * 25.4) * mm,
        bulletIndent=style.bulletIndent + depth * (_INDENT / 72 * 25.4) * mm,
    )


def _table_as_lines(rows: richtext.Grid | list[list[str]]) -> list[str]:
    """A table as one `기준 · A사 · B사` line per row, for the `.hwpx` fallback."""
    kept = [row for row in _as_grid(rows).flat() if any(cell.strip() for cell in row)]
    return [" · ".join(_strip_inline(cell) for cell in row) for row in kept]


def _pdf_chart(chart: dict, style: dict | None):
    """The chart as a reportlab drawing sized to the text column. The value axis starts at zero."""
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.charts.legends import Legend
    from reportlab.graphics.charts.linecharts import HorizontalLineChart
    from reportlab.graphics.shapes import Drawing, String
    from reportlab.graphics.widgets.markers import makeMarker

    accent = HexColor(style["accent"]) if style else HexColor("#5b5bd6")
    muted = HexColor(style["muted"]) if style else HexColor("#666666")
    korean = fonts.korean(style["font"] if style else "serif")
    width, height = 170 * mm, 80 * mm

    drawing = Drawing(width, height)
    plot = (HorizontalLineChart if chart["kind"] == "line" else VerticalBarChart)()
    plot.x, plot.y = 22 * mm, 14 * mm
    plot.width, plot.height = width - 30 * mm, height - 24 * mm
    plot.data = [list(values) for _name, values in chart["series"]]
    plot.categoryAxis.categoryNames = list(chart["categories"])
    plot.categoryAxis.labels.fontName = korean
    plot.categoryAxis.labels.fontSize = 8
    plot.categoryAxis.labels.fillColor = muted
    plot.valueAxis.valueMin = 0
    plot.valueAxis.labels.fontName = korean
    plot.valueAxis.labels.fontSize = 8
    plot.valueAxis.labels.fillColor = muted

    shades = [accent, _lighter(accent, 0.55)]
    for index in range(len(plot.data)):
        colour = shades[index % len(shades)]
        if chart["kind"] == "line":
            plot.lines[index].strokeColor = colour
            plot.lines[index].strokeWidth = 1.6
            # Markers, as on the Word chart.
            plot.lines[index].symbol = makeMarker("FilledCircle")
            plot.lines[index].symbol.fillColor = colour
            plot.lines[index].symbol.strokeColor = colour
            plot.lines[index].symbol.size = 4
        else:
            plot.bars[index].fillColor = colour
            plot.bars[index].strokeColor = None
    drawing.add(plot)

    if unit := chart.get("unit"):
        drawing.add(String(4, height - 12, unit, fontName=korean, fontSize=8, fillColor=muted))
    # A legend only for more than one series.
    if len(chart["series"]) > 1:
        legend = Legend()
        legend.x, legend.y = 22 * mm, 4 * mm
        legend.alignment = "right"
        legend.fontName, legend.fontSize = korean, 8
        legend.dxTextSpace, legend.deltax = 4, 60
        legend.columnMaximum = 1
        legend.colorNamePairs = [
            (shades[index % len(shades)], name or "값")
            for index, (name, _values) in enumerate(chart["series"])
        ]
        drawing.add(legend)
    return drawing


def _lighter(colour, amount: float):
    """One colour mixed toward white — a second series, from one accent."""
    from reportlab.lib.colors import Color

    return Color(
        colour.red + (1 - colour.red) * amount,
        colour.green + (1 - colour.green) * amount,
        colour.blue + (1 - colour.blue) * amount,
    )


def _pdf_cards(
    cards: list[tuple[str, list[str]]], styles: dict, style: dict | None
) -> Table | None:
    """The card grid, drawn. Two columns of boxes, the title in accent."""
    if not cards:
        return None
    accent = HexColor(style["accent"]) if style else HexColor("#5b5bd6")
    hair = HexColor(style["muted"]) if style else HexColor("#666666")
    head = ParagraphStyle(
        "cardTitle",
        parent=styles["body"],
        fontSize=11,
        leading=15,
        alignment=TA_LEFT,
        textColor=accent,
        spaceAfter=3,
    )
    item = ParagraphStyle(
        "cardItem",
        parent=styles["body"],
        fontSize=9.5,
        leading=13.5,
        alignment=TA_LEFT,
        spaceAfter=1,
    )

    def cell(card: tuple[str, list[str]] | None) -> list:
        if card is None:
            return []
        title, items = card
        return [Paragraph(f"<b>{_escape(title)}</b>", head)] + [
            Paragraph(f"· {_escape(line)}", item) for line in items
        ]

    padded: list[tuple[str, list[str]] | None] = list(cards)
    if len(padded) % 2:
        # The odd card out gets an empty cell.
        padded.append(None)
    rows = [[cell(padded[i]), cell(padded[i + 1])] for i in range(0, len(padded), 2)]
    table = Table(rows, colWidths=[85 * mm, 85 * mm])
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("BOX", (0, 0), (-1, -1), 0.4, hair),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, hair),
            ]
        )
    )
    return table


def _pdf_callout(callout: tuple[str, list[str]], styles: dict, style: dict | None) -> Table | None:
    """One box, with a bar down its left edge — the shape of a thing not to skip."""
    title, lines = callout
    if not title:
        return None
    accent = HexColor(style["accent"]) if style else HexColor("#5b5bd6")
    head = ParagraphStyle(
        "calloutTitle",
        parent=styles["body"],
        fontSize=11,
        leading=15,
        alignment=TA_LEFT,
        textColor=accent,
        spaceAfter=2,
    )
    body = ParagraphStyle(
        "calloutBody", parent=styles["body"], fontSize=10, leading=15, alignment=TA_LEFT
    )
    inner = [Paragraph(f"<b>{_escape(title)}</b>", head)] + [
        Paragraph(_escape(line), body) for line in lines
    ]
    table = Table([[inner]], colWidths=[170 * mm])
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                # A rule down the left rather than a box.
                ("LINEBEFORE", (0, 0), (0, -1), 2.2, accent),
            ]
        )
    )
    return table


def _pdf_steps(steps: list[tuple[str, str]], styles: dict, style: dict | None) -> Table | None:
    """The same procedure, drawn. A rail down the left, not a grid."""
    if not steps:
        return None
    accent = HexColor(style["accent"]) if style else HexColor("#5b5bd6")
    muted = HexColor(style["muted"]) if style else HexColor("#666666")
    number = ParagraphStyle(
        "stepNumber",
        parent=styles["body"],
        fontSize=12,
        leading=15,
        alignment=TA_CENTER,
        textColor=accent,
    )
    rows = [
        [
            Paragraph(str(i + 1), number),
            Paragraph(
                f"<b>{_escape(name)}</b>" + (f"  {_escape(detail)}" if detail else ""),
                styles["body"],
            ),
        ]
        for i, (name, detail) in enumerate(steps)
    ]
    table = Table(rows, colWidths=[12 * mm, 158 * mm])
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (1, 0), (1, -1), 6),
                ("LINEAFTER", (0, 0), (0, -1), 0.8, accent),
                ("LINEBELOW", (0, 0), (-1, -2), 0.4, muted),
            ]
        )
    )
    return table


def _pdf_kpi(figures: list[tuple[str, str]], styles: dict, style: dict | None) -> Table | None:
    """The same strip, drawn. Values above, labels under, no rules."""
    if not figures:
        return None
    accent = HexColor(style["accent"]) if style else HexColor("#5b5bd6")
    muted = HexColor(style["muted"]) if style else HexColor("#666666")
    big = ParagraphStyle(
        "kpiValue",
        parent=styles["body"],
        fontSize=doc_type.TYPE["kpi"],
        leading=doc_type.TYPE["kpi"] * 1.2,
        alignment=TA_CENTER,
        textColor=accent,
    )
    small = ParagraphStyle(
        "kpiLabel",
        parent=styles["body"],
        fontSize=doc_type.TYPE["kpiLabel"],
        leading=doc_type.TYPE["kpiLabel"] * 1.3,
        alignment=TA_CENTER,
        textColor=muted,
    )
    table = Table(
        [
            [Paragraph(_escape(value), big) for value, _ in figures],
            [Paragraph(_escape(label), small) for _, label in figures],
        ],
        colWidths=[(170 * mm) / len(figures)] * len(figures),
    )
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, 0), 2),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 1),
                ("TOPPADDING", (0, 1), (-1, 1), 0),
                ("BOTTOMPADDING", (0, 1), (-1, 1), 2),
                ("LINEABOVE", (0, 0), (-1, 0), 0.6, muted),
                ("LINEBELOW", (0, 1), (-1, 1), 0.6, muted),
            ]
        )
    )
    return table


def _pdf_picture(picture: dict, styles: dict) -> list:
    """One figure as flowables, or `[]` when it cannot be drawn."""
    data = picture.get("data")
    if not data:
        return []
    try:
        width_pt, height_pt = _picture_size(data)
        out: list = [
            Spacer(1, 3 * mm),
            RLImage(io.BytesIO(data), width=width_pt, height=height_pt),
        ]
    except Exception as exc:  # noqa: BLE001 — a bad picture is not a failed export
        log.warning("could not place a picture in the pdf: %s", exc)
        return []
    if caption := str(picture.get("caption") or ""):
        out.append(Paragraph(_escape(caption), styles["caption"]))
    out.append(Spacer(1, 3 * mm))
    return out


def _pdf_table(rows: richtext.Grid | list[list[str]], styles: dict, accent) -> Table | None:
    """A drawn table, or `None` when there is nothing in it.

    Cells are `Paragraph`s so long text wraps; merges are `SPAN` commands over
    the covered rectangle.
    """
    grid = richtext.Grid(
        rows=[row for row in _as_grid(rows).rows if any(c.text.strip() for c in row)]
    )
    if not grid.rows:
        return None
    width = grid.width
    body = [[Paragraph("", styles["body"]) for _ in range(width)] for _ in grid.rows]
    spans: list[tuple] = []
    covered: set[tuple[int, int]] = set()
    for r, row in enumerate(grid.rows):
        column = 0
        for source in row:
            while (r, column) in covered:
                column += 1
            if column >= width:
                break
            across = min(source.colspan, width - column)
            down = min(source.rowspan, len(grid.rows) - r)
            for dr in range(down):
                for dc in range(across):
                    if (dr, dc) != (0, 0):
                        covered.add((r + dr, column + dc))
            if across > 1 or down > 1:
                spans.append(("SPAN", (column, r), (column + across - 1, r + down - 1)))
            text = "<br/>".join(_escape(_strip_inline(line)) for line in source.text.split("\n"))
            body[r][column] = Paragraph(text, styles["body"])
            column += across

    table = Table(body, colWidths=[(170 * mm) / width] * width, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                *spans,
                ("GRID", (0, 0), (-1, -1), 0.4, HexColor("#d0d0d0")),
                ("BACKGROUND", (0, 0), (-1, 0), HexColor("#f2f2f2")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LINEBELOW", (0, 0), (-1, 0), 1.0, accent),
            ]
        )
    )
    return table


def to_pdf(
    title: str,
    sections: list[dict],
    *,
    tokens: dict[str, str] | None = None,
    page_settings: dict | None = None,
) -> bytes:
    # Embedded Korean face — see services/fonts.py. Gothic unless the design says serif,
    # as on the page view; the sizes are the document scale every renderer reads.
    style = design.normalise_tokens(tokens) if tokens else None
    korean = fonts.korean(style["font"] if style else "gothic")
    bold_face = fonts.korean(style["font"] if style else "gothic", bold=True)
    T, L = doc_type.TYPE, doc_type.LEADING
    # Without a design system the headings stay black.
    heading_colour = {"textColor": HexColor(style["accent"])} if style else {}

    visual_style = (style or {}).get("visualStyle") or "editorial"
    accent_colour = HexColor(style["accent"]) if style else HexColor("#5b5bd6")
    base = getSampleStyleSheet()
    title_size, h1_size = T["title"], T["h1"]
    title_options = (
        {
            "fontSize": title_size,
            "leading": title_size * L["title"],
            "textColor": HexColor("#ffffff"),
            "backColor": accent_colour,
            "borderPadding": 18,
        }
        if visual_style == "poster"
        else {
            "fontSize": title_size,
            "leading": title_size * L["title"],
            "textColor": HexColor(style["ink"] if style else "#1a1a1a"),
        }
        if visual_style == "minimal"
        else {"fontSize": title_size, "leading": title_size * L["title"], **heading_colour}
    )
    h1_options = (
        {
            "fontSize": h1_size,
            "leading": h1_size * L["heading"],
            "textColor": accent_colour,
            "spaceBefore": 18,
            "spaceAfter": 9,
        }
        if visual_style == "poster"
        else {
            "fontSize": h1_size,
            "leading": h1_size * L["heading"],
            "textColor": HexColor(style["muted"] if style else "#666666"),
            "spaceBefore": 16,
            "spaceAfter": 5,
        }
        if visual_style == "minimal"
        else {
            "fontSize": h1_size,
            "leading": h1_size * L["heading"],
            "spaceBefore": 14,
            "spaceAfter": 6,
            **heading_colour,
        }
    )
    styles = {
        "title": ParagraphStyle(
            "t",
            parent=base["Title"],
            fontName=bold_face,
            **title_options,
        ),
        "h1": ParagraphStyle(
            "h1",
            parent=base["Heading1"],
            fontName=bold_face,
            **h1_options,
        ),
        "h2": ParagraphStyle(
            "h2",
            parent=base["Heading2"],
            fontName=bold_face,
            fontSize=T["h2"],
            leading=T["h2"] * L["heading"],
            spaceBefore=10,
            spaceAfter=4,
        ),
        # Left-aligned: Korean has no hyphenation, and justified lines with long
        # Latin tokens gap.
        "body": ParagraphStyle(
            "b",
            parent=base["BodyText"],
            fontName=korean,
            fontSize=T["body"],
            leading=T["body"] * L["body"],
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "li",
            parent=base["BodyText"],
            fontName=korean,
            fontSize=T["body"],
            leading=T["body"] * L["body"],
            leftIndent=10 * mm,
            bulletIndent=4 * mm,
            spaceAfter=3,
        ),
        # Section footnotes: small, hanging under their mark.
        "note": ParagraphStyle(
            "note",
            parent=base["BodyText"],
            fontName=korean,
            fontSize=T["note"],
            leading=T["note"] * L["note"],
            textColor=HexColor(style["muted"]) if style else HexColor("#666666"),
            leftIndent=6 * mm,
            firstLineIndent=-6 * mm,
            spaceAfter=1,
        ),
        # Centred under a centred picture.
        "caption": ParagraphStyle(
            "cap",
            parent=base["BodyText"],
            fontName=korean,
            fontSize=T["caption"],
            leading=T["caption"] * L["note"],
            alignment=TA_CENTER,
            textColor=HexColor(style["muted"]) if style else HexColor("#666666"),
            spaceBefore=2,
            spaceAfter=2,
        ),
    }

    story: list = [Paragraph(_escape(title), styles["title"]), Spacer(1, 8 * mm)]
    for index, section in enumerate(sections):
        if index:
            story.append(Spacer(1, 4 * mm))
        story.append(Paragraph(_escape(section.get("heading") or ""), styles["h1"]))
        #: This section's footnotes, drawn under it once the prose is done.
        notes: list[tuple[str, str]] = []
        formatted = list(section.get("_formatting") or [])
        format_cursor = 0
        for kind, text, marker, depth in _markdown_to_lines(
            section.get("content") or "", section.get("tables")
        ):
            if kind == "pagebreak":
                story.append(PageBreak())
                continue
            if kind == "table":
                accent = HexColor(style["accent"]) if style else HexColor("#5b5bd6")
                if drawn := _pdf_table(text, styles, accent):
                    story.append(Spacer(1, 2 * mm))
                    story.append(drawn)
                    story.append(Spacer(1, 3 * mm))
                continue
            if kind == "image":
                for flowable in _pdf_picture(text, styles):
                    story.append(flowable)
                continue
            if kind == "note":
                # Gathered under the section: reportlab has no footnote model.
                notes.append((marker, _strip_inline(text)))
                continue
            if kind == "chart":
                story.append(Spacer(1, 3 * mm))
                story.append(_pdf_chart(text, style))
                story.append(Spacer(1, 4 * mm))
                continue
            if kind == "steps":
                if drawn := _pdf_steps(text, styles, style):
                    story.append(Spacer(1, 2 * mm))
                    story.append(drawn)
                    story.append(Spacer(1, 4 * mm))
                continue
            if kind == "kpi":
                if strip := _pdf_kpi(text, styles, style):
                    strip = KeepTogether(strip)
                    story.append(Spacer(1, 3 * mm))
                    story.append(strip)
                    story.append(Spacer(1, 4 * mm))
                continue
            if kind == "cards":
                if grid := _pdf_cards(text, styles, style):
                    story.append(Spacer(1, 3 * mm))
                    story.append(grid)
                    story.append(Spacer(1, 4 * mm))
                continue
            if kind == "callout":
                if boxed := _pdf_callout(text, styles, style):
                    story.append(Spacer(1, 3 * mm))
                    story.append(boxed)
                    story.append(Spacer(1, 4 * mm))
                continue
            if kind == "diagram":
                if picture := _diagram_picture(section, text):
                    for flowable in _pdf_picture(picture, styles):
                        story.append(flowable)
                else:
                    story.append(
                        Paragraph(
                            _escape("[다이어그램 — 화면에서 한 번 열면 그림으로 저장됩니다]"),
                            styles["caption"],
                        )
                    )
                continue
            plain = _strip_inline(text)
            block, format_cursor = _next_format(formatted, plain, format_cursor)
            clean = _MARK.sub(
                # reportlab's Paragraph understands `<super>`.
                lambda found: f"<super>{_escape(found.group(1))}</super>",
                _escape(plain),
            )
            if block and not _MARK.search(plain):
                clean = _pdf_styled(block, plain)
            if kind == "heading":
                paragraph_style = _pdf_block_style(styles["h2"], block) if block else styles["h2"]
                story.append(Paragraph(clean, paragraph_style))
            elif kind in ("bullet", "number"):
                # `bulletText` hangs the marker, keeping two-digit numbers
                # aligned with single-digit ones.
                story.append(
                    Paragraph(
                        clean,
                        _pdf_block_style(_at_depth(styles["bullet"], depth), block)
                        if block
                        else _at_depth(styles["bullet"], depth),
                        bulletText=marker,
                    )
                )
            else:
                paragraph_style = (
                    _pdf_block_style(styles["body"], block) if block else styles["body"]
                )
                story.append(Paragraph(clean, paragraph_style))
        for picture in section.get("images") or []:
            data = picture.get("data")
            if not data:
                continue
            width, height = _picture_size(data)
            caption = str(picture.get("caption") or "")
            try:
                image = RLImage(io.BytesIO(data), width=width, height=height)
                image.hAlign = "CENTER"
                figure: list = [image]
            except Exception as exc:  # noqa: BLE001 — a bad picture is not a failed export
                log.warning("could not place a picture in the report pdf: %s", exc)
                continue
            if caption:
                figure.append(Paragraph(_escape(caption), styles["caption"]))
            story.append(Spacer(1, 3 * mm))
            # Kept together so the caption stays on the picture's page.
            story.append(KeepTogether(figure))
            story.append(Spacer(1, 3 * mm))

        if notes:
            # A rule, then the notes.
            story.append(Spacer(1, 3 * mm))
            story.append(
                HRFlowable(
                    width="35%",
                    thickness=0.5,
                    color=HexColor(style["muted"]) if style else HexColor("#666666"),
                    spaceAfter=2,
                    hAlign="LEFT",
                )
            )
            for mark, note in notes:
                story.append(
                    Paragraph(f"<super>{_escape(mark)}</super> {_escape(note)}", styles["note"])
                )

    buffer = io.BytesIO()
    page_settings = page_settings or {}
    margins = page_settings.get("margins") if isinstance(page_settings.get("margins"), dict) else {}

    def decorate(pdf, document, page: int, total: int) -> None:
        pdf.saveState()
        pdf.setFont(korean, 8)
        pdf.setFillColor(HexColor("#777777"))
        header = str(page_settings.get("header") or title)
        if header and (page > 1 or bool(page_settings.get("firstPageHeader", False))):
            pdf.drawString(document.leftMargin, A4[1] - 10 * mm, header)
        footer = str(page_settings.get("footer") or "")
        if footer:
            pdf.drawString(document.leftMargin, 9 * mm, footer)
        numbering = str(page_settings.get("pageNumbers") or "page-total")
        if numbering != "none":
            label = f"{page} / {total}" if numbering == "page-total" else str(page)
            pdf.drawRightString(A4[0] - document.rightMargin, 9 * mm, label)
        pdf.restoreState()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=float(margins.get("left", 22)) * mm,
        rightMargin=float(margins.get("right", 22)) * mm,
        topMargin=float(margins.get("top", 22)) * mm,
        bottomMargin=float(margins.get("bottom", 20)) * mm,
        title=title,
    )

    class NumberedCanvas(pdfcanvas.Canvas):
        """Replays every page once the total is known, so `1 / 12` is right without building twice."""

        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._page_states: list[dict] = []

        def showPage(self) -> None:  # noqa: N802 - ReportLab API
            self._page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self) -> None:
            total = len(self._page_states)
            for page, state in enumerate(self._page_states, 1):
                self.__dict__.update(state)
                decorate(self, document, page, total)
                super().showPage()
            super().save()

    document.build(story, canvasmaker=NumberedCanvas)
    return buffer.getvalue()


#: The text column width, and the most height one figure may take, in mm.
_PICTURE_MM = 150.0
_PICTURE_MAX_MM = 170.0

#: Pixels are read at 96 DPI, the rate Hancom uses.
_POINTS_PER_PIXEL = 72 / 96


def _picture_size(data: bytes) -> tuple[float, float]:
    """`(width, height)` in points: native size, shrunk only if it overflows the column."""
    try:
        with PIL.Image.open(io.BytesIO(data)) as picture:
            pixels_wide, pixels_high = picture.size
    except Exception:  # noqa: BLE001 — an unreadable picture still gets a box
        pixels_wide, pixels_high = 480, 320
    width = max(1, pixels_wide) * _POINTS_PER_PIXEL
    height = max(1, pixels_high) * _POINTS_PER_PIXEL
    scale = min(1.0, _PICTURE_MM * mm / width, _PICTURE_MAX_MM * mm / height)
    return width * scale, height * scale


def _escape(text: str) -> str:
    """reportlab's Paragraph reads its input as mini-HTML."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ── HWPX (OWPML, KS X 6101) ───────────────────────────────────────────
#
# An XML zip like .docx, written with the standard library alone. Hancom
# Office's reader is stricter than Word's: a missing part, or a mimetype entry
# in the wrong position, is rejected rather than repaired.

_HWPX_MIMETYPE = "application/hwp+zip"

_HWPX_VERSION = """<?xml version="1.0" encoding="UTF-8"?>
<hv:HCFVersion xmlns:hv="http://www.hancom.co.kr/hwpml/2011/version" tagetApplication="WORDPROCESSOR"
 major="5" minor="1" micro="1" buildNumber="0" os="1" xmlVersion="1.4" application="KloudChat"/>"""

_HWPX_CONTAINER = """<?xml version="1.0" encoding="UTF-8"?>
<ocf:container xmlns:ocf="urn:oasis:names:tc:opendocument:xmlns:container"
 xmlns:hpf="http://www.hancom.co.kr/schema/2011/hpf">
 <ocf:rootfiles>
  <ocf:rootfile full-path="Contents/content.hpf" media-type="application/hwpml-package+xml"/>
 </ocf:rootfiles>
</ocf:container>"""

_HWPX_MANIFEST = """<?xml version="1.0" encoding="UTF-8"?>
<odf:manifest xmlns:odf="urn:oasis:names:tc:opendocument:xmlns:manifest">
 <odf:file-entry odf:full-path="/" odf:media-type="application/hwp+zip"/>
 <odf:file-entry odf:full-path="Contents/header.xml" odf:media-type="application/xml"/>
 <odf:file-entry odf:full-path="Contents/section0.xml" odf:media-type="application/xml"/>
</odf:manifest>"""

_HWPX_CONTENT_HPF = """<?xml version="1.0" encoding="UTF-8"?>
<hpf:package xmlns:hpf="http://www.hancom.co.kr/schema/2011/hpf"
 xmlns:opf="http://www.idpf.org/2007/opf/" version="" unique-identifier="" id="">
 <opf:metadata><opf:title>{title}</opf:title></opf:metadata>
 <opf:manifest>
  <opf:item id="header" href="Contents/header.xml" media-type="application/xml"/>
  <opf:item id="section0" href="Contents/section0.xml" media-type="application/xml"/>
 </opf:manifest>
 <opf:spine><opf:itemref idref="header"/><opf:itemref idref="section0"/></opf:spine>
</hpf:package>"""

#: section0.xml refers to character and paragraph shapes by index, so reordering
#: corrupts the body. `<hh:lineSpacing>` is mandatory (absent, Hancom reads 0%),
#: and alignment must be the `<hh:align>` child element: the HWPML 2010
#: `paraPr@align` attribute is silently ignored.
_HWPX_HEADER = """<?xml version="1.0" encoding="UTF-8"?>
<hh:head xmlns:hh="http://www.hancom.co.kr/hwpml/2011/head"
 xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core" version="1.4" secCnt="1">
 <hh:beginNum page="1" footnote="1" endnote="1" pic="1" tbl="1" equation="1"/>
 <hh:refList>
  <hh:fontfaces itemCnt="1">
   <hh:fontface lang="HANGUL" fontCnt="1">
    <hh:font id="0" face="함초롬바탕" type="TTF" isEmbedded="0"/>
   </hh:fontface>
  </hh:fontfaces>
{border_fills}
{char_properties}
{para_properties}
  <hh:styles itemCnt="1">
   <hh:style id="0" type="PARA" name="바탕글" engName="Normal" paraPrIDRef="3" charPrIDRef="0" nextStyleIDRef="0" langID="1042"/>
  </hh:styles>
 </hh:refList>
</hh:head>"""

#: Border definitions a cell points at by `borderFillIDRef`; an undefined id
#: stops the file opening. Must sit between `fontfaces` and `charProperties`.
#: Ids are 1-based: 1 none (the table frame), 2 body cell hairline, 3 head cell
#: with a heavier line under it.
_HWPX_BORDER_FILLS = """  <hh:borderFills itemCnt="3">
   <hh:borderFill id="1" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0">
    <hh:slash type="NONE" Crooked="0" isCounter="0"/>
    <hh:backSlash type="NONE" Crooked="0" isCounter="0"/>
    <hh:leftBorder type="NONE" width="0.1 mm" color="#000000"/>
    <hh:rightBorder type="NONE" width="0.1 mm" color="#000000"/>
    <hh:topBorder type="NONE" width="0.1 mm" color="#000000"/>
    <hh:bottomBorder type="NONE" width="0.1 mm" color="#000000"/>
    <hh:diagonal type="NONE" width="0.1 mm" color="#000000"/>
   </hh:borderFill>
   <hh:borderFill id="2" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0">
    <hh:slash type="NONE" Crooked="0" isCounter="0"/>
    <hh:backSlash type="NONE" Crooked="0" isCounter="0"/>
    <hh:leftBorder type="SOLID" width="0.1 mm" color="#BBBBBB"/>
    <hh:rightBorder type="SOLID" width="0.1 mm" color="#BBBBBB"/>
    <hh:topBorder type="SOLID" width="0.1 mm" color="#BBBBBB"/>
    <hh:bottomBorder type="SOLID" width="0.1 mm" color="#BBBBBB"/>
    <hh:diagonal type="NONE" width="0.1 mm" color="#000000"/>
   </hh:borderFill>
   <hh:borderFill id="3" threeD="0" shadow="0" centerLine="NONE" breakCellSeparateLine="0">
    <hh:slash type="NONE" Crooked="0" isCounter="0"/>
    <hh:backSlash type="NONE" Crooked="0" isCounter="0"/>
    <hh:leftBorder type="SOLID" width="0.1 mm" color="#BBBBBB"/>
    <hh:rightBorder type="SOLID" width="0.1 mm" color="#BBBBBB"/>
    <hh:topBorder type="SOLID" width="0.1 mm" color="#BBBBBB"/>
    <!-- 굵은 아래선만. fillBrush 머리행 음영은 한컴이 무시한다. -->
    <hh:bottomBorder type="SOLID" width="0.5 mm" color="#444444"/>
    <hh:diagonal type="NONE" width="0.1 mm" color="#000000"/>
   </hh:borderFill>
  </hh:borderFills>"""

#: The text column, in HWPUNIT: page width less both margins. A table wider
#: than this is a table Hancom pushes off the page.
_HWPX_TEXT_WIDTH = 59528 - 8504 - 8504

#: A 10pt full-width Hangul character in HWPUNIT; column widths are reasoned about in characters.
_HWPX_BODY_CHAR = 1000
#: A row's nominal height; Hancom grows cells to fit.
_HWPX_ROW_HEIGHT = 2000
#: Padding inside a cell, all four sides.
_HWPX_CELL_MARGIN = 141


#: `charPr@height` is in 1/100 pt.
_HWPX_CHAR_SHAPES = (
    # (id, height in 1/100 pt, bold) — the document scale, as the page view and the .docx.
    (0, doc_type.hwp("body"), False),  # body
    (1, doc_type.hwp("body"), True),  # body, bold
    (2, doc_type.hwp("title"), True),  # document title
    (3, doc_type.hwp("h1"), True),  # section heading  (h1)
    (4, doc_type.hwp("h2"), True),  # sub-heading      (h2)
    (5, doc_type.hwp("note"), False),  # footnote — smaller than the prose, and not bold
)

#: (id, horizontal align, left indent, space-before, space-after) in HWPUNIT
#: (1/7200 in, so 1000 == 10 pt).
_HWPX_PARA_SHAPES = (
    (0, "CENTER", 0, 0, 600),  # title
    (1, "LEFT", 0, 600, 300),  # h1
    (2, "LEFT", 0, 400, 200),  # h2
    (3, "JUSTIFY", 0, 0, 150),  # body
    (4, "JUSTIFY", 1000, 0, 100),  # bullet — indented from the body margin
    (5, "CENTER", 0, 300, 100),  # figure and its caption
    (6, "CENTER", 0, 0, 0),  # a centred table cell — no space of its own
    # A table cell, left-aligned: justified Hangul in a narrow cell is pulled apart.
    (7, "LEFT", 0, 0, 0),
    # Sub-bullet levels. Appended at the end: ids are positions and tests read by index.
    (8, "JUSTIFY", 2000, 0, 100),  # bullet, one level in
    (9, "JUSTIFY", 3000, 0, 100),  # bullet, two levels in
)

#: 160% is the usual line spacing for a Korean report.
_HWPX_LINE_SPACING = 160


#: Character-shape ids the accent colours: the title and section headings only.
_HWPX_ACCENT_SHAPES = (2, 3)


def _hwpx_char_properties(
    accent: str | None = None,
    extras: list[tuple[int, int, bool, bool, bool, bool, str, str]] | None = None,
    visual_style: str = "editorial",
) -> str:
    items = []
    shapes = [
        (
            cid,
            height,
            bold,
            False,
            False,
            False,
            accent if (accent and cid in _HWPX_ACCENT_SHAPES) else "#000000",
            "none",
        )
        for cid, height, bold in _HWPX_CHAR_SHAPES
    ]
    if visual_style == "poster":
        shapes = [
            (
                cid,
                2400 if cid == 2 else 1600 if cid == 3 else height,
                bold,
                italic,
                underline,
                strike,
                "#FFFFFF" if cid == 2 else colour,
                accent if cid == 2 and accent else shade,
            )
            for cid, height, bold, italic, underline, strike, colour, shade in shapes
        ]
    elif visual_style == "minimal":
        shapes = [
            (
                cid,
                1800 if cid == 2 else 1100 if cid == 3 else height,
                bold if cid != 3 else False,
                italic,
                underline,
                strike,
                "#666666" if cid == 3 else colour,
                shade,
            )
            for cid, height, bold, italic, underline, strike, colour, shade in shapes
        ]
    shapes += list(extras or [])
    for cid, height, bold, italic, underline, strike, colour, shade in shapes:
        items.append(
            f'   <hh:charPr id="{cid}" height="{height}" textColor="{colour}"'
            f' shadeColor="{shade}" useFontSpace="0" useKerning="0">\n'
            '    <hh:fontRef hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>\n'
            '    <hh:ratio hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/>\n'
            '    <hh:spacing hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>\n'
            '    <hh:relSz hangul="100" latin="100" hanja="100" japanese="100" other="100" symbol="100" user="100"/>\n'
            '    <hh:offset hangul="0" latin="0" hanja="0" japanese="0" other="0" symbol="0" user="0"/>\n'
            + ("    <hh:bold/>\n" if bold else "")
            + ("    <hh:italic/>\n" if italic else "")
            + (
                '    <hh:underline type="BOTTOM" shape="SOLID" color="#000000"/>\n'
                if underline
                else ""
            )
            + ('    <hh:strikeout shape="SOLID" color="#000000"/>\n' if strike else "")
            + "   </hh:charPr>"
        )
    return (
        f'  <hh:charProperties itemCnt="{len(items)}">\n'
        + "\n".join(items)
        + "\n  </hh:charProperties>"
    )


def _hwpx_para_properties(
    extras: list[tuple[int, str, int, int, int, int]] | None = None,
) -> str:
    items = []
    shapes = [(*shape, _HWPX_LINE_SPACING) for shape in _HWPX_PARA_SHAPES] + list(extras or [])
    for pid, align, left, prev, nxt, line_spacing in shapes:
        items.append(
            f'   <hh:paraPr id="{pid}" tabPrIDRef="0" condense="0" fontLineHeight="0"'
            ' snapToGrid="1" suppressLineNumbers="0" checked="0">\n'
            f'    <hh:align horizontal="{align}" vertical="BASELINE"/>\n'
            '    <hh:heading type="NONE" idRef="0" level="0"/>\n'
            '    <hh:breakSetting breakLatinWord="KEEP_WORD" breakNonLatinWord="KEEP_WORD"'
            ' widowOrphan="0" keepWithNext="0" keepLines="0" pageBreakBefore="0"'
            ' lineWrap="BREAK"/>\n'
            '    <hh:autoSpacing eAsianEng="0" eAsianNum="0"/>\n'
            "    <hh:margin>\n"
            '     <hc:intent value="0" unit="HWPUNIT"/>\n'
            f'     <hc:left value="{left}" unit="HWPUNIT"/>\n'
            '     <hc:right value="0" unit="HWPUNIT"/>\n'
            f'     <hc:prev value="{prev}" unit="HWPUNIT"/>\n'
            f'     <hc:next value="{nxt}" unit="HWPUNIT"/>\n'
            "    </hh:margin>\n"
            f'    <hh:lineSpacing type="PERCENT" value="{line_spacing}" unit="HWPUNIT"/>\n'
            "   </hh:paraPr>"
        )
    return (
        f'  <hh:paraProperties itemCnt="{len(items)}">\n'
        + "\n".join(items)
        + "\n  </hh:paraProperties>"
    )


def _hwpx_escape(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _hwpx_secpr(page_settings: dict | None = None) -> str:
    """OWPML section properties: A4 (59528 x 84188 HWPUNIT), margins, header and page-number visibility.

    Carried in the first paragraph's run. Without a page box an absolutely sized
    picture is read and not drawn.
    """
    settings = page_settings or {}
    margins = settings.get("margins") if isinstance(settings.get("margins"), dict) else {}

    def unit(name: str, fallback: float) -> int:
        try:
            value = max(10.0, min(35.0, float(margins.get(name, fallback))))
        except (TypeError, ValueError):
            value = fallback
        return round(value * 7200 / 25.4)

    first_header = "0" if bool(settings.get("firstPageHeader", False)) else "1"
    hide_page_number = "1" if str(settings.get("pageNumbers") or "page-total") == "none" else "0"
    return (
        '<hp:secPr id="" textDirection="HORIZONTAL" spaceColumns="1134" tabStop="8000"'
        ' tabStopVal="4000" tabStopUnit="HWPUNIT" outlineShapeIDRef="0" memoShapeIDRef="0"'
        ' textVerticalWidthHead="0" masterPageCnt="0">'
        '<hp:grid lineGrid="0" charGrid="0" wonggojiFormat="0"/>'
        '<hp:startNum pageStartsOn="BOTH" page="0" pic="0" tbl="0" equation="0"/>'
        f'<hp:visibility hideFirstHeader="{first_header}" hideFirstFooter="0" hideFirstMasterPage="0"'
        f' border="SHOW_ALL" fill="SHOW_ALL" hideFirstPageNum="{hide_page_number}" hideFirstEmptyLine="0"'
        ' showLineNumber="0"/>'
        '<hp:pagePr landscape="WIDELY" width="59528" height="84188" gutterType="LEFT_ONLY">'
        f'<hp:margin header="{unit("top", 20)}" footer="{unit("bottom", 15)}" gutter="0"'
        f' left="{unit("left", 30)}" right="{unit("right", 30)}"'
        f' top="{unit("top", 20)}" bottom="{unit("bottom", 15)}"/></hp:pagePr>'
        "</hp:secPr>"
    )


#: HWPUNIT is 1/7200 inch, and Hancom reads a picture's pixels at 96 DPI
#: whatever the file's own metadata says: one pixel is 75 HWPUNIT.
_HWPUNIT_PER_PIXEL = 75

#: The text column of the page above, and as much height as a figure may take
#: before it owns the page.
_HWPX_MAX_WIDTH = 59528 - 8504 * 2
_HWPX_MAX_HEIGHT = int(170 / 25.4 * 7200)

#: One inline picture. `binaryItemIDRef` resolves against the `<opf:item>` id in
#: `Contents/content.hpf`; nothing is declared in `header.xml`.
_HWPX_PIC = (
    '<hp:pic id="{n}" zOrder="0" numberingType="PICTURE" textWrap="SQUARE"'
    ' textFlow="BOTH_SIDES" lock="0" dropcapstyle="None" href="" groupLevel="0"'
    ' instid="{n}" reverse="0">'
    '<hp:offset x="0" y="0"/>'
    '<hp:orgSz width="{w}" height="{h}"/>'
    '<hp:curSz width="{w}" height="{h}"/>'
    '<hp:flip horizontal="0" vertical="0"/>'
    '<hp:rotationInfo angle="0" centerX="{cx}" centerY="{cy}" rotateimage="1"/>'
    "<hp:renderingInfo>"
    '<hc:transMatrix e1="1" e2="0" e3="0" e4="0" e5="1" e6="0"/>'
    '<hc:scaMatrix e1="1" e2="0" e3="0" e4="0" e5="1" e6="0"/>'
    '<hc:rotMatrix e1="1" e2="0" e3="0" e4="0" e5="1" e6="0"/>'
    "</hp:renderingInfo>"
    '<hp:imgRect><hc:pt0 x="0" y="0"/><hc:pt1 x="{w}" y="0"/>'
    '<hc:pt2 x="{w}" y="{h}"/><hc:pt3 x="0" y="{h}"/></hp:imgRect>'
    '<hp:imgClip left="0" right="{dw}" top="0" bottom="{dh}"/>'
    '<hp:inMargin left="0" right="0" top="0" bottom="0"/>'
    '<hp:imgDim dimwidth="{dw}" dimheight="{dh}"/>'
    '<hc:img binaryItemIDRef="{ref}" bright="0" contrast="0" effect="REAL_PIC" alpha="0"/>'
    "<hp:effects/>"
    '<hp:sz width="{w}" widthRelTo="ABSOLUTE" height="{h}" heightRelTo="ABSOLUTE" protect="0"/>'
    '<hp:pos treatAsChar="1" affectLSpacing="0" flowWithText="1" allowOverlap="0"'
    ' holdAnchorAndSO="0" vertRelTo="PARA" horzRelTo="PARA" vertAlign="TOP"'
    ' horzAlign="LEFT" vertOffset="0" horzOffset="0"/>'
    '<hp:outMargin left="0" right="0" top="0" bottom="0"/>'
    "</hp:pic>"
)


def _hwpx_picture(index: int, data: bytes) -> str:
    """One picture paragraph, sized to the page.

    `imgDim` and `imgClip` stay in the picture's own pixels — they are the
    source rectangle — while `orgSz`, `curSz`, `sz` and `imgRect` carry the
    size on the page. They are equal for a picture that already fits.
    """
    try:
        with PIL.Image.open(io.BytesIO(data)) as picture:
            pixels_wide, pixels_high = picture.size
    except Exception:  # noqa: BLE001 — an unreadable picture gets a plain box
        pixels_wide, pixels_high = 480, 320
    native_w = max(1, pixels_wide) * _HWPUNIT_PER_PIXEL
    native_h = max(1, pixels_high) * _HWPUNIT_PER_PIXEL
    scale = min(1.0, _HWPX_MAX_WIDTH / native_w, _HWPX_MAX_HEIGHT / native_h)
    width, height = int(native_w * scale), int(native_h * scale)
    return (
        '<hp:p paraPrIDRef="5" styleIDRef="0"><hp:run charPrIDRef="0">'
        + _HWPX_PIC.format(
            n=index,
            ref=f"image{index}",
            w=width,
            h=height,
            dw=native_w,
            dh=native_h,
            cx=width // 2,
            cy=height // 2,
        )
        + "<hp:t/></hp:run></hp:p>"
    )


def _hwpx_para(text: str, para_pr: int, char_pr: int = 0) -> str:
    """One `<hp:p>`. An empty run still needs the `<hp:t/>` — Hancom renders a
    paragraph with no run as a missing line rather than a blank one."""
    return (
        f'<hp:p paraPrIDRef="{para_pr}" styleIDRef="0">'
        f'<hp:run charPrIDRef="{char_pr}">'
        f"<hp:t>{_hwpx_escape(text)}</hp:t>"
        f"</hp:run></hp:p>"
    )


def _hwpx_page_furniture(title: str, page_settings: dict | None) -> str:
    """Header, footer and live page-number controls for the section's first run."""
    if not page_settings:
        return ""
    header = str(page_settings.get("header") or title).strip()
    footer = str(page_settings.get("footer") or "").strip()

    def control(kind: str, text: str, identifier: int) -> str:
        if not text:
            return ""
        paragraph = _hwpx_para(text, 3, 5)
        return (
            f'<hp:ctrl><hp:{kind} id="{identifier}" applyPageType="BOTH">'
            '<hp:subList id="" textDirection="HORIZONTAL" lineWrap="BREAK"'
            ' vertAlign="TOP" linkListIDRef="0" linkListNextIDRef="0"'
            ' textWidth="0" textHeight="0" hasTextRef="0" hasNumRef="0">'
            f"{paragraph}</hp:subList></hp:{kind}></hp:ctrl>"
        )

    numbering = str(page_settings.get("pageNumbers") or "page-total")
    page_number = (
        '<hp:ctrl><hp:pageNum pos="BOTTOM_RIGHT" formatType="DIGIT" sideChar=""/></hp:ctrl>'
        if numbering != "none"
        else ""
    )
    return control("header", header, 10001) + control("footer", footer, 10002) + page_number


def _column_weights(rows: list[list[str]], cols: int) -> list[int]:
    """Column widths in characters, summing to what the text column holds.

    Each column first gets its longest word (Hancom splits Hangul between
    characters with no hyphen), then the rest is shared by text length. When
    the words alone do not fit, short needs are met first and the shortfall
    falls on the columns that can wrap.
    """

    def longest_word(column: int) -> int:
        return max(
            (len(word) for row in rows if column < len(row) for word in row[column].split()),
            default=1,
        )

    #: Characters the text column holds once cell margins are out.
    budget = (_HWPX_TEXT_WIDTH - cols * _HWPX_CELL_MARGIN * 2) // _HWPX_BODY_CHAR
    need = [max(2, min(longest_word(c), budget // 2)) for c in range(cols)]
    if sum(need) >= budget:
        # Short needs met first, in ascending order; the shortfall lands on the wide columns.
        out = [0] * cols
        rest = budget
        for taken, column in enumerate(sorted(range(cols), key=lambda c: need[c])):
            share = rest // (cols - taken)
            out[column] = min(need[column], share)
            rest -= out[column]
        return out
    want = [max((len(row[c]) for row in rows if c < len(row)), default=1) for c in range(cols)]
    spare = budget - sum(need)
    extra = [max(0, w - n) for w, n in zip(want, need, strict=True)]
    total = sum(extra)
    if total == 0:
        return need
    out = [n + spare * e // total for n, e in zip(need, extra, strict=True)]
    out[-1] += budget - sum(out)
    return out


def _hwpx_table(
    rows: richtext.Grid | list[list[str]],
    *,
    char_pr: int = 0,
    head_char_pr: int = 1,
    widths: list[int] | None = None,
    cell_para_pr: int | list[int] = 7,
) -> str:
    """A GFM table as an OWPML table, wrapped in its own paragraph.

    Child order is load-bearing: `sz`, `pos`, `outMargin`, `inMargin`, rows
    inside `hp:tbl`; `subList`, `cellAddr`, `cellSpan`, `cellSz`, `cellMargin`
    inside `hp:tc`. Every `borderFillIDRef` must exist in `_HWPX_BORDER_FILLS`.
    `widths` are relative column weights; `cell_para_pr` is 7 (left) or 6
    (centred), or one per column. Cells covered by a merge are omitted, as OWPML
    expects. Raises on anything it cannot render so the caller falls back to lines.
    """
    grid = richtext.Grid(
        rows=[row for row in _as_grid(rows).rows if any(c.text.strip() for c in row)]
    )
    if not grid.rows:
        raise ValueError("빈 표")
    kept = grid.flat()
    cols = grid.width
    if cols < 1 or cols > 12:
        raise ValueError(f"열이 {cols}개")
    # Sums to exactly the text width; a rounding remainder shows as a short column.
    weights = (widths or _column_weights(kept, cols))[:cols]
    weights += [1] * (cols - len(weights))
    total = sum(weights) or cols
    column_widths = [_HWPX_TEXT_WIDTH * w // total for w in weights]
    column_widths[-1] += _HWPX_TEXT_WIDTH - sum(column_widths)

    body: list[str] = []
    #: Addresses covered by a merge that began above or to the left. OWPML wants
    #: them absent, not empty.
    covered: set[tuple[int, int]] = set()
    for r, row in enumerate(grid.rows):
        cells: list[str] = []
        column = 0
        for cell in row:
            while (r, column) in covered:
                column += 1
            if column >= cols:
                break
            across = min(cell.colspan, cols - column)
            down = min(cell.rowspan, len(grid.rows) - r)
            for dr in range(down):
                for dc in range(across):
                    if (dr, dc) != (0, 0):
                        covered.add((r + dr, column + dc))
            # Head row: bold at body size with a heavier rule under it.
            fill = 3 if r == 0 else 2
            shape = cell_para_pr[column] if isinstance(cell_para_pr, list) else cell_para_pr
            lines = [_strip_inline(line) for line in cell.text.split("\n")] or [""]
            para = "".join(
                _hwpx_para(line, shape, head_char_pr if r == 0 else char_pr) for line in lines
            )
            width = sum(column_widths[column : column + across])
            cells.append(
                f'<hp:tc name="" header="{1 if r == 0 else 0}" hasMargin="1" protect="0"'
                f' editable="0" dirty="0" borderFillIDRef="{fill}">'
                '<hp:subList id="" textDirection="HORIZONTAL" lineWrap="BREAK"'
                ' vertAlign="TOP" linkListIDRef="0" linkListNextIDRef="0" textWidth="0"'
                ' textHeight="0" hasTextRef="0" hasNumRef="0">'
                f"{para}"
                "</hp:subList>"
                f'<hp:cellAddr colAddr="{column}" rowAddr="{r}"/>'
                f'<hp:cellSpan colSpan="{across}" rowSpan="{down}"/>'
                f'<hp:cellSz width="{width}" height="{_HWPX_ROW_HEIGHT * down}"/>'
                f'<hp:cellMargin left="{_HWPX_CELL_MARGIN}" right="{_HWPX_CELL_MARGIN}"'
                f' top="{_HWPX_CELL_MARGIN}" bottom="{_HWPX_CELL_MARGIN}"/>'
                "</hp:tc>"
            )
            column += across
        body.append("<hp:tr>" + "".join(cells) + "</hp:tr>")

    height = _HWPX_ROW_HEIGHT * len(grid.rows)
    table = (
        f'<hp:tbl id="0" zOrder="0" numberingType="TABLE" textWrap="TOP_AND_BOTTOM"'
        ' textFlow="BOTH_SIDES" lock="0" dropcapstyle="None" pageBreak="CELL"'
        f' repeatHeader="1" rowCnt="{len(grid.rows)}" colCnt="{cols}" cellSpacing="0"'
        ' borderFillIDRef="1" noAdjust="0">'
        f'<hp:sz width="{_HWPX_TEXT_WIDTH}" widthRelTo="ABSOLUTE" height="{height}"'
        ' heightRelTo="ABSOLUTE" protect="0"/>'
        '<hp:pos treatAsChar="1" affectLSpacing="0" flowWithText="1" allowOverlap="0"'
        ' holdAnchorAndSO="0" vertRelTo="PARA" horzRelTo="COLUMN" vertAlign="TOP"'
        ' horzAlign="LEFT" vertOffset="0" horzOffset="0"/>'
        '<hp:outMargin left="0" right="0" top="0" bottom="283"/>'
        f'<hp:inMargin left="{_HWPX_CELL_MARGIN}" right="{_HWPX_CELL_MARGIN}"'
        f' top="{_HWPX_CELL_MARGIN}" bottom="{_HWPX_CELL_MARGIN}"/>' + "".join(body) + "</hp:tbl>"
    )
    # `treatAsChar` puts the table in the run, so it gets its own paragraph.
    return f'<hp:p paraPrIDRef="3" styleIDRef="0"><hp:run charPrIDRef="{char_pr}">{table}</hp:run></hp:p>'


def to_hwpx(
    title: str,
    sections: list[dict],
    *,
    tokens: dict[str, str] | None = None,
    page_settings: dict | None = None,
) -> bytes:
    """The same document `to_docx` writes, as OWPML.

    Bullets are literal `•` text: a wrong numbering reference makes Hancom
    refuse the file. Pictures are `BinData/` parts declared in `content.hpf`
    and the OCF manifest; any mismatch is a file that will not open. A diagram
    with no raster yet is written as `[다이어그램]`.
    """
    style = design.normalise_tokens(tokens) if tokens else None
    char_extras: list[tuple[int, int, bool, bool, bool, bool, str, str]] = []
    char_ids: dict[tuple, int] = {}
    para_extras: list[tuple[int, str, int, int, int, int]] = []
    para_ids: dict[tuple, int] = {}

    def colour(value: str, fallback: str) -> str:
        if found := re.fullmatch(r"#([0-9a-f]{6})", value, re.I):
            return "#" + found.group(1).upper()
        if found := re.fullmatch(r"#([0-9a-f]{3})", value, re.I):
            return "#" + "".join(character * 2 for character in found.group(1)).upper()
        return fallback

    def styled_char_id(run_style: dict, base: int) -> int:
        _, base_height, base_bold = _HWPX_CHAR_SHAPES[base]
        size = str(run_style.get("font-size") or "")
        height = (
            round(float(found.group(1)) * 100)
            if (found := re.fullmatch(r"(\d+(?:\.\d+)?)pt", size, re.I))
            else base_height
        )
        weight = str(run_style.get("font-weight") or "").lower()
        decoration = str(run_style.get("text-decoration") or "").lower()
        key = (
            max(600, min(7200, height)),
            base_bold or weight in {"bold", "600", "700", "800", "900"},
            str(run_style.get("font-style") or "").lower() == "italic",
            "underline" in decoration,
            "line-through" in decoration,
            colour(str(run_style.get("color") or ""), "#000000"),
            colour(str(run_style.get("background-color") or ""), "none"),
        )
        if key not in char_ids:
            identifier = len(_HWPX_CHAR_SHAPES) + len(char_extras)
            char_ids[key] = identifier
            char_extras.append((identifier, *key))
        return char_ids[key]

    def styled_para_id(block: dict, base: int) -> int:
        _, default_align, left, prev, nxt = _HWPX_PARA_SHAPES[base]
        block_style = block.get("style") or {}
        align = {
            "left": "LEFT",
            "center": "CENTER",
            "right": "RIGHT",
            "justify": "JUSTIFY",
        }.get(str(block_style.get("text-align") or "").lower(), default_align)
        try:
            spacing = round(float(block_style.get("line-height") or "") * 100)
        except (TypeError, ValueError):
            spacing = _HWPX_LINE_SPACING
        key = (align, left, prev, nxt, max(100, min(300, spacing)))
        if key == (default_align, left, prev, nxt, _HWPX_LINE_SPACING):
            return base
        if key not in para_ids:
            identifier = len(_HWPX_PARA_SHAPES) + len(para_extras)
            para_ids[key] = identifier
            para_extras.append((identifier, *key))
        return para_ids[key]

    def styled_para(block: dict, base_para: int, base_char: int, prefix: str = "") -> str:
        runs: list[str] = []
        if prefix:
            runs.append(
                f'<hp:run charPrIDRef="{base_char}"><hp:t>{_hwpx_escape(prefix)}</hp:t></hp:run>'
            )
        for item in block.get("runs") or []:
            text = _hwpx_escape(str(item.get("text") or "")).replace("\n", "<hp:lineBreak/>")
            runs.append(
                f'<hp:run charPrIDRef="{styled_char_id(item.get("style") or {}, base_char)}">'
                f"<hp:t>{text}</hp:t></hp:run>"
            )
        return (
            f'<hp:p paraPrIDRef="{styled_para_id(block, base_para)}" styleIDRef="0">'
            + "".join(runs)
            + "</hp:p>"
        )

    # The section properties ride in the first paragraph's run.
    body: list[str] = [
        f'<hp:p paraPrIDRef="0" styleIDRef="0"><hp:run charPrIDRef="2">'
        f"{_hwpx_secpr(page_settings)}{_hwpx_page_furniture(title, page_settings)}"
        f"<hp:t>{_hwpx_escape(title)}</hp:t></hp:run></hp:p>"
    ]
    #: `BinData/imageN.png` and the `<opf:item id="imageN">` that resolves it.
    embedded: list[tuple[str, bytes, str]] = []

    def place(picture: dict) -> list[str]:
        """One picture embedded where it stands, with its caption under it."""
        data = picture.get("data")
        if not data:
            return []
        mime = str(picture.get("mime") or "image/png").lower()
        # `image/jpg` is Hancom's own spelling; the extension follows it.
        extension = {"image/jpeg": "jpg", "image/gif": "gif", "image/webp": "webp"}.get(mime, "png")
        index = len(embedded) + 1
        embedded.append((f"image{index}", data, extension))
        out = [_hwpx_picture(index, data)]
        caption = str(picture.get("caption") or "").strip()
        if caption:
            out.append(_hwpx_para(caption, 5, 4))
        return out

    for section in sections:
        heading = (section.get("heading") or "").strip()
        if heading:
            body.append(_hwpx_para(heading, 1, 3))
        formatted = list(section.get("_formatting") or [])
        format_cursor = 0
        #: Footnotes gathered under the section; `<hp:footNote>` is not used
        #: because a malformed one will not open.
        notes: list[tuple[str, str]] = []
        for kind, text, marker, depth in _markdown_to_lines(
            section.get("content") or "", section.get("tables")
        ):
            if kind == "pagebreak":
                body.append(
                    '<hp:p paraPrIDRef="3" styleIDRef="0" pageBreak="1">'
                    '<hp:run charPrIDRef="0"><hp:t/></hp:run></hp:p>'
                )
                continue
            if kind == "note":
                notes.append((marker, _strip_inline(text)))
                continue
            if kind == "table":
                try:
                    body.append(_hwpx_table(text))
                except Exception as exc:  # noqa: BLE001 — see `_hwpx_table`
                    log.warning("hwpx table fell back to lines: %s", exc)
                    for line in _table_as_lines(text):
                        body.append(_hwpx_para(line, 3))
                continue
            if kind == "image":
                body.extend(place(text))
                continue
            if kind == "chart":
                # The browser's raster if there is one, else the numbers as a
                # table. OWPML's chart element is not written.
                if drawn := _diagram_picture(section, {"key": diagram_key(marker)}):
                    body.extend(place(drawn))
                    continue
                try:
                    body.append(_hwpx_table(_chart_as_rows(text), head_char_pr=2))
                except Exception as exc:  # noqa: BLE001 — see `_hwpx_table`
                    log.warning("hwpx chart fell back to lines: %s", exc)
                    for line in _table_as_lines(_chart_as_rows(text)):
                        body.append(_hwpx_para(line, 3))
                continue
            if kind == "steps":
                try:
                    body.append(
                        _hwpx_table(
                            [["", "단계", "내용"]]
                            + [[str(i + 1), name, detail] for i, (name, detail) in enumerate(text)],
                            head_char_pr=2,
                            widths=[1, 5, 12],
                            cell_para_pr=[6, 7, 7],
                        )
                    )
                except Exception as exc:  # noqa: BLE001 — see `_hwpx_table`
                    log.warning("hwpx steps fell back to lines: %s", exc)
                    for i, (name, detail) in enumerate(text):
                        body.append(_hwpx_para(f"{i + 1}. {name} {detail}".strip(), 3))
                continue
            if kind == "cards":
                # Two columns; a card is its title and lines in one cell.
                try:
                    grid = [
                        [
                            "\n".join([title] + [f"· {line}" for line in items])
                            for title, items in pair
                        ]
                        for pair in _in_pairs(text)
                    ]
                    body.append(_hwpx_table(grid, cell_para_pr=3))
                except Exception as exc:  # noqa: BLE001 — see `_hwpx_table`
                    log.warning("hwpx cards fell back to lines: %s", exc)
                    for title, items in text:
                        body.append(_hwpx_para(title, 2))
                        for line in items:
                            body.append(_hwpx_para(f"· {line}", 3))
                continue
            if kind == "callout":
                # One cell: 한글 has no left rule that survives a re-save.
                title, lines = text
                try:
                    body.append(_hwpx_table([["\n".join([title] + list(lines))]], cell_para_pr=3))
                except Exception as exc:  # noqa: BLE001 — see `_hwpx_table`
                    log.warning("hwpx callout fell back to lines: %s", exc)
                    body.append(_hwpx_para(title, 2))
                    for line in lines:
                        body.append(_hwpx_para(line, 3))
                continue
            if kind == "kpi":
                try:
                    body.append(
                        _hwpx_table(
                            [[v for v, _ in text], [label for _, label in text]],
                            head_char_pr=2,
                            cell_para_pr=6,
                        )
                    )
                except Exception as exc:  # noqa: BLE001 — see `_hwpx_table`
                    log.warning("hwpx kpi fell back to lines: %s", exc)
                    body.append(_hwpx_para(" · ".join(f"{v} {label}" for v, label in text), 3))
                continue
            if kind == "diagram":
                if drawn := _diagram_picture(section, text):
                    body.extend(place(drawn))
                else:
                    body.append(_hwpx_para("[다이어그램]", 5, 0))
                continue
            clean = _raised_marks(_strip_inline(text))
            block, format_cursor = _next_format(formatted, _strip_inline(text), format_cursor)
            if kind == "heading":
                body.append(styled_para(block, 2, 4) if block else _hwpx_para(clean, 2, 4))
            elif kind in ("bullet", "number"):
                base_para = (4, 8, 9)[depth]
                body.append(
                    styled_para(block, base_para, 0, f"{marker} ")
                    if block
                    else _hwpx_para(f"{marker} {clean}", base_para)
                )
            else:
                body.append(styled_para(block, 3, 0) if block else _hwpx_para(clean, 3))
        for picture in section.get("images") or []:
            body.extend(place(picture))
        for mark, note in notes:
            body.append(_hwpx_para(f"{_raised_marks(f'[^{mark}]')} {note}", 3, 5))

    section_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<hs:sec xmlns:hs="http://www.hancom.co.kr/hwpml/2011/section"'
        ' xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph"'
        ' xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core">' + "".join(body) + "</hs:sec>"
    )

    # PrvText is what a file manager previews.
    preview = "\n".join([title] + [(s.get("heading") or "") for s in sections])[:1000]

    # One `<opf:item>` per picture; `isEmbeded` (OWPML's spelling) stops Hancom dropping it.
    items = "".join(
        f'  <opf:item id="{name}" href="BinData/{name}.{extension}"'
        f' media-type="image/{"jpg" if extension == "jpg" else extension}" isEmbeded="1"/>\n'
        for name, _, extension in embedded
    )
    content_hpf = _HWPX_CONTENT_HPF.format(title=_hwpx_escape(title)).replace(
        '  <opf:item id="section0"', items + '  <opf:item id="section0"', 1
    )
    manifest = _HWPX_MANIFEST.replace(
        "</odf:manifest>",
        "".join(
            f' <odf:file-entry odf:full-path="BinData/{name}.{extension}"'
            f' odf:media-type="image/{"jpg" if extension == "jpg" else extension}"/>\n'
            for name, _, extension in embedded
        )
        + "</odf:manifest>",
        1,
    )

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        # `mimetype` must be first and STORED, as in ODF/EPUB.
        archive.writestr(
            zipfile.ZipInfo("mimetype"), _HWPX_MIMETYPE, compress_type=zipfile.ZIP_STORED
        )
        archive.writestr("version.xml", _HWPX_VERSION)
        archive.writestr("META-INF/container.xml", _HWPX_CONTAINER)
        archive.writestr("META-INF/manifest.xml", manifest)
        archive.writestr("Contents/content.hpf", content_hpf)
        archive.writestr(
            "Contents/header.xml",
            _HWPX_HEADER.format(
                border_fills=_HWPX_BORDER_FILLS,
                char_properties=_hwpx_char_properties(
                    style["accent"] if style else None,
                    char_extras,
                    (style or {}).get("visualStyle") or "editorial",
                ),
                para_properties=_hwpx_para_properties(para_extras),
            ),
        )
        archive.writestr("Contents/section0.xml", section_xml)
        for name, data, extension in embedded:
            # Binaries are STORED, as Hancom writes them.
            archive.writestr(
                zipfile.ZipInfo(f"BinData/{name}.{extension}"),
                data,
                compress_type=zipfile.ZIP_STORED,
            )
        archive.writestr("Preview/PrvText.txt", preview)
    return buffer.getvalue()


__all__ = ["to_docx", "to_pdf", "to_hwpx"]
