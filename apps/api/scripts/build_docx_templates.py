"""Generates one real Word template per document 서식.

Run from the API image:

    docker compose run --rm --no-deps api python scripts/build_docx_templates.py

The problem this solves: `to_docx` wrote the same generic Word document
whichever 서식 the person picked. The 서식 shaped the page view and the printed
HTML, and then the file — the thing that actually gets submitted — came out in
`python-docx`'s defaults. Somebody who chose 회의록 and downloaded a `.docx`
got a document with none of 회의록 about it.

A template fixes that at the root rather than in the writer. `python-docx`
opens a `.docx` and inherits its styles, its page setup and its theme, so
`Document(<서식>/template.docx)` *is* the 서식 — the writer goes on calling
`add_heading` and `add_paragraph`, and what those mean is now the template's.

Generated rather than authored because a `.docx` is a zip of XML: there is no
way to hand-write one and no reason to check a binary into the tree without a
recipe beside it. The recipe is `_SPECS` below, and it is deliberately close to
what the matching `seed.html` declares — the two are the same design, one for
screen and one for Word, and where they disagree the document changes shape
depending on which one somebody opened.

Fonts name Korean faces every Hangul Windows and Hancom install carries. A
template that names a face the reader does not have is a template Word
silently substitutes, which is worse than not setting one.
"""

from __future__ import annotations

import pathlib
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent / "app" / "design_templates"


class Spec:
    """One 서식's Word half, in the terms the seed states its own design in."""

    def __init__(
        self,
        folder: str,
        *,
        body_font: str,
        heading_font: str,
        accent: str,
        body_pt: float,
        margins_mm: tuple[int, int, int, int],
        title_pt: float,
        h1_pt: float,
        h2_pt: float,
        note: str,
        form: tuple[tuple, ...] = (),
    ) -> None:
        self.folder = folder
        self.form = form
        self.body_font = body_font
        self.heading_font = heading_font
        self.accent = accent
        self.body_pt = body_pt
        self.margins_mm = margins_mm
        self.title_pt = title_pt
        self.h1_pt = h1_pt
        self.h2_pt = h2_pt
        self.note = note


# ── the shape of the blank form ───────────────────────────────────────────
#
# `template.docx` is styles and page only, because the writer appends to it and
# anything left inside would arrive at the top of every document written from
# it. So the form somebody downloads is a second file built from the same
# styles: `form.docx`, with the headings and the tables in place and a line of
# guidance under each.
#
# Written as data rather than as code per 서식 so the two halves cannot drift:
# one spec makes the style base the writer uses and the form the reader fills
# in by hand, and a heading that exists in one exists in the other.


def H(level: int, text: str) -> tuple:
    """A heading. Level 0 is the document's title."""
    return ("h", level, text)


def P(text: str) -> tuple:
    """A line of guidance, greyed: what belongs here, not an example of it."""
    return ("p", text)


def T(columns: tuple[str, ...], rows: int = 3) -> tuple:
    """A table with its header filled in and `rows` empty rows under it."""
    return ("t", columns, rows)


def B(items: tuple[str, ...]) -> tuple:
    """A bulleted list of prompts, one per line the reader is expected to add."""
    return ("b", items)


#: Kept beside the seeds on purpose. When a seed's type scale changes, this is
#: the other half of that change — a document that reads one way on screen and
#: another in Word is two documents wearing one name.
_SPECS = (
    Spec(
        "doc-report",
        body_font="함초롬바탕",
        heading_font="함초롬돋움",
        accent="2B4C7E",
        body_pt=10.5,
        margins_mm=(20, 20, 20, 20),
        title_pt=20,
        h1_pt=13,
        h2_pt=11,
        note="제출용 보고 문서. 각주로 근거를 단다.",
        form=(
            H(0, "보고서 제목"),
            P("표지만 읽고 무엇을 위한 문서인지, 누가 읽는지, 언제 것인지 알 수 있게."),
            T(("작성", "소속", "작성일", "문서번호"), rows=1),
            H(1, "검토 범위"),
            P("무엇을 보고 무엇을 보지 않았는지. 범위를 밝히지 않으면 빠진 것이 실수로 읽힙니다."),
            H(1, "확인한 것"),
            P(
                "수치와 인용에는 각주 표시를 답니다. "
                "표시 없는 수치가 이어지면 근거를 적지 않은 것입니다."
            ),
            T(("항목", "값", "단위·기준 시점", "출처")),
            H(1, "비교"),
            P("같은 기준으로 견줍니다. 비교 값이라면 무엇 대비인지 밝힙니다."),
            H(1, "권고와 남은 확인"),
            P("확인하지 못한 것은 확인하지 못했다고 적습니다."),
            H(1, "읽는 사람이 할 일"),
            T(("할 일", "담당", "기한")),
        ),
    ),
    Spec(
        "doc-project-brief",
        body_font="함초롬돋움",
        heading_font="함초롬돋움",
        accent="3B5BDB",
        body_pt=10,
        margins_mm=(18, 18, 18, 18),
        title_pt=19,
        h1_pt=12.5,
        h2_pt=10.5,
        note="프로젝트 브리프. 표지 아래 개요 한 문단, 나머지는 나란히 놓는 격자.",
        form=(
            H(0, "프로젝트 제목"),
            P("이 프로젝트가 무엇을 하는 것인지 한 줄로."),
            T(("책임자", "기간", "작성일"), rows=1),
            H(1, "개요"),
            P("무엇을, 왜, 언제까지. 한 문단으로 씁니다. 세 문단이면 개요가 아닙니다."),
            H(1, "산출물과 목표"),
            P("받는 사람이 손에 쥐는 것과, 확인할 수 있는 기준. 활동은 산출물이 아닙니다."),
            T(("산출물", "목표")),
            H(1, "이해관계자와 성공 기준"),
            P("누가 무엇을 하는지와, 무엇을 보면 됐다고 할 수 있는지."),
            T(("이해관계자", "역할", "성공 기준")),
            H(1, "일정과 예산"),
            T(("단계", "기간", "금액")),
            H(1, "전제와 위험"),
            B(("틀리면 무엇이 무너지는가", "무엇을 보면 아는가")),
        ),
    ),
    Spec(
        "doc-brief",
        body_font="함초롬돋움",
        heading_font="함초롬돋움",
        accent="1F6F5C",
        body_pt=10,
        margins_mm=(16, 16, 16, 16),
        title_pt=17,
        h1_pt=11.5,
        h2_pt=10.5,
        note="한 장 요약. 여백을 줄여 한 쪽에 담는다.",
        form=(
            H(0, "한 장 요약 제목"),
            P("읽는 사람이 무엇을 결정해야 하는지 한 줄로."),
            T(("작성", "작성일", "결정 기한"), rows=1),
            H(1, "결정할 것"),
            P("첫 줄에서 무엇을 정해야 하는지 알 수 있게."),
            H(1, "지금 상황"),
            P("숫자로. 형용사만 있으면 그 문장은 의견입니다."),
            H(1, "대안"),
            P("같은 기준으로 비교합니다. 하나만 있으면 비교가 아닙니다."),
            T(("안", "장점", "단점", "비용")),
            H(1, "걸려 있는 것"),
            B(("무엇을 잃는가", "무엇을 얻는가")),
            H(1, "다음 행동"),
            P("조판이 마지막 항목을 강조하므로 다른 것을 여기 두지 않습니다."),
            T(("할 일", "담당", "기한")),
        ),
    ),
    Spec(
        "doc-minutes",
        body_font="함초롬바탕",
        heading_font="함초롬돋움",
        accent="5B4B8A",
        body_pt=10.5,
        margins_mm=(20, 18, 18, 18),
        title_pt=18,
        h1_pt=12.5,
        h2_pt=11,
        note="회의록. 결정과 논의를 갈라 적는다.",
        form=(
            H(0, "회의 이름"),
            T(("일시", "장소", "작성"), rows=1),
            H(1, "참석"),
            P("오지 못한 사람이 이 문서만 읽고 따라올 수 있게. 불참자도 적습니다."),
            H(1, "논의"),
            P("정해진 것과 논의만 된 것을 섞지 않습니다. 반대나 유보 의견도 남깁니다."),
            H(1, "결정 사항"),
            P("결정마다 왜 그렇게 정했는지 한 줄이라도 남깁니다."),
            T(("결정", "근거", "담당", "기한")),
            H(1, "남은 쟁점"),
            P("결론이 나지 않은 안건은 다음에 누가 언제 정하는지 적습니다."),
            H(1, "실행 항목"),
            P("오른쪽 칸은 문장이 아니라 이름과 날짜만. 미정이면 '미정' 이라고 적습니다."),
            T(("할 일", "담당", "기한")),
        ),
    ),
    Spec(
        "doc-notice",
        body_font="함초롬돋움",
        heading_font="함초롬돋움",
        accent="A2452C",
        body_pt=11,
        margins_mm=(22, 20, 20, 20),
        title_pt=19,
        h1_pt=13,
        h2_pt=11.5,
        note="안내문·공지. 본문을 크게 잡아 게시용으로 읽힌다.",
        form=(
            T(("수신", "참조"), rows=1),
            H(0, "제목"),
            T(("시행일", "담당", "연락처", "문서번호"), rows=1),
            H(1, "근거"),
            P("이 안내가 무엇에 근거한 것인지. 규정·회의 결정·상위 공문 등."),
            H(1, "내용"),
            P("첫 절을 읽고 무엇이 언제부터 달라지는지 알 수 있게."),
            B(("무엇이", "언제부터", "누구에게 해당하는지")),
            H(1, "협조 요청"),
            P("받는 사람이 할 일에 기한과 제출처를 함께 적습니다. 할 일이 없으면 없다고 말합니다."),
            H(1, "붙임"),
            P("본문에서 실제로 쓰인 자료만. 붙임 1. 파일 이름 1부.  끝."),
        ),
    ),
    Spec(
        "doc-lab",
        body_font="함초롬바탕",
        heading_font="함초롬돋움",
        accent="1F5673",
        body_pt=10,
        margins_mm=(18, 18, 18, 18),
        title_pt=17,
        h1_pt=12,
        h2_pt=10.5,
        note="실험 노트. 절차와 관찰을 촘촘히 담는다.",
        form=(
            H(0, "실험 제목"),
            P("무엇을 확인하려 했는지 표지에서 알 수 있게."),
            T(("실험자", "실험일", "장소", "지도"), rows=1),
            H(1, "목적"),
            P("무엇을 확인하려는 실험인지 한 문장으로."),
            H(1, "장비·설정값·반복 횟수"),
            P("이 노트만 보고 같은 실험을 다시 할 수 있게."),
            T(("장비", "설정값", "반복")),
            H(1, "절차"),
            B(("준비", "측정", "정리")),
            H(1, "측정값"),
            P("첫 칸만 항목 이름, 나머지는 수치. 단위는 열 제목에만 답니다."),
            T(("회차", "측정값", "조건")),
            H(1, "계산"),
            P(
                "결과 숫자가 어떤 식에서 나왔는지 전개를 남깁니다. "
                "값만 있고 계산이 없으면 안 됩니다."
            ),
            H(1, "해석"),
            P(
                "관측한 값과 그에 대한 해석을 구분합니다. "
                "예상과 다른 값은 지우지 않고 그대로 둡니다."
            ),
            H(1, "한계"),
            P("실제 한계를 적습니다. '추가 연구가 필요하다' 로 대신하지 않습니다."),
            H(1, "원자료"),
            P("원자료가 어디 있는지."),
        ),
    ),
    Spec(
        "doc-term-paper",
        body_font="함초롬바탕",
        heading_font="함초롬돋움",
        accent="4B3B8F",
        body_pt=10.5,
        margins_mm=(25, 25, 25, 25),
        title_pt=19,
        h1_pt=13,
        h2_pt=11,
        note="기말 리포트. 각주가 앉을 여백을 넓게 둔다.",
        form=(
            H(0, "리포트 제목"),
            P("무엇을 묻는 글인지 부제로 한 줄."),
            T(("과목", "학번", "이름", "제출일"), rows=1),
            H(1, "서론"),
            P("이 글이 답하려는 질문을 문장으로 세웁니다. 배경만 늘어놓지 않습니다."),
            H(1, "본론"),
            P("장 제목이 그 장의 주장이 되게 씁니다. 인용한 문장에는 그 자리에 각주를 답니다."),
            H(1, "결론"),
            P("서론에서 세운 질문에 답합니다. 앞을 요약만 하지 않습니다."),
            H(1, "다루지 못한 것"),
            P("이 글이 다루지 못한 것을 한 줄로 밝힙니다."),
            H(1, "참고문헌"),
            P("각주에 단 것과 어긋나지 않게 모읍니다."),
        ),
    ),
    Spec(
        "doc-case",
        body_font="함초롬돋움",
        heading_font="함초롬돋움",
        accent="1F6FEB",
        body_pt=10.5,
        margins_mm=(20, 18, 18, 18),
        title_pt=19,
        h1_pt=13,
        h2_pt=11,
        note="케이스 분석. 표로 견주고 하나를 고른다.",
        form=(
            H(0, "케이스 제목"),
            T(("대상 기업", "분석 기간", "작성", "작성일"), rows=1),
            H(1, "권고"),
            P("무엇을 하자는 것인지 먼저 적습니다. 뒤의 절이 이것의 근거가 됩니다."),
            H(1, "현황"),
            P("숫자로 씁니다. 형용사만 있으면 그 문장은 의견입니다."),
            T(("지표", "기준 시점", "값", "출처")),
            H(1, "대안"),
            P("열은 대안, 행은 기준. 같은 줄에서 견줄 수 있어야 합니다."),
            T(("기준", "안 1", "안 2", "안 3")),
            H(1, "권고"),
            P("어느 안을, 어떤 기준에서, 무엇을 포기하고 고르는지."),
            H(1, "위험"),
            B(("틀렸을 때 무엇이 일어나는가", "무엇을 보면 틀린 줄 아는가")),
        ),
    ),
    Spec(
        "doc-survey",
        body_font="함초롬돋움",
        heading_font="함초롬돋움",
        accent="0F766E",
        body_pt=10,
        margins_mm=(20, 18, 18, 18),
        title_pt=18,
        h1_pt=12.5,
        h2_pt=11,
        note="설문 분석. 표가 많으므로 본문을 작게.",
        form=(
            H(0, "조사 제목"),
            T(("조사 기간", "대상", "표본 수", "작성"), rows=1),
            H(1, "조사 개요"),
            P("누구에게, 언제, 어떻게 물었는지."),
            T(("배포", "회수", "회수율"), rows=1),
            H(1, "표본"),
            P("치우쳤다면 치우쳤다고 본문에 적습니다."),
            T(("구분", "빈도", "비율(%)")),
            H(1, "기술통계"),
            P("표는 숫자만. 해석은 표 아래 문단에."),
            T(("문항", "N", "평균", "표준편차")),
            H(1, "교차분석"),
            P("표 제목에 걸어 본 두 변수를 적습니다. 검정을 했다면 이름과 값도."),
            H(1, "한계"),
            B(("표본의 치우침", "자기보고의 한계", "묻지 못한 것")),
        ),
    ),
    Spec(
        "doc-incident",
        body_font="함초롬돋움",
        heading_font="함초롬돋움",
        accent="B42318",
        body_pt=10.5,
        margins_mm=(18, 18, 18, 18),
        title_pt=18,
        h1_pt=12.5,
        h2_pt=11,
        note="장애 보고. 시간순 기록이 넓게 앉는다.",
        form=(
            H(0, "장애 제목"),
            T(("발생", "인지", "복구", "영향 범위"), rows=1),
            H(1, "요약"),
            P("무엇이, 얼마 동안, 누구에게. 세 문장 안에."),
            H(1, "시간순 기록"),
            P("추측과 사실을 섞지 않습니다. 그때 몰랐던 것은 몰랐다고 적습니다."),
            T(("시각(KST)", "일어난 일", "한 일")),
            H(1, "원인"),
            P("왜 그 조작이 가능했는지, 왜 막히지 않았는지, 왜 늦게 알았는지까지."),
            H(1, "재발 방지"),
            T(("할 일", "담당(역할)", "기한")),
            H(1, "탐지"),
            P("늦게 알았다면 무엇을 보면 빨리 알 수 있었는지."),
        ),
    ),
    Spec(
        "doc-proposal",
        body_font="함초롬돋움",
        heading_font="함초롬돋움",
        accent="7C3AED",
        body_pt=10.5,
        margins_mm=(20, 20, 20, 20),
        title_pt=20,
        h1_pt=13,
        h2_pt=11,
        note="제안서. 표지를 크게 잡아 보내는 문서로 읽힌다.",
        form=(
            H(0, "제안 제목"),
            P("무엇을 제안하는지 한 줄로."),
            T(("고객사", "제안일", "담당", "연락처"), rows=1),
            H(1, "고객의 과제"),
            P("우리가 아니라 고객의 말로 시작합니다."),
            H(1, "제안"),
            P("무엇을 하겠다는 것인지 한 문장으로 먼저."),
            H(1, "효과"),
            P("무엇이 얼마나, 어떤 계산으로. 계산의 전제를 밝힙니다."),
            T(("항목", "지금", "도입 후", "근거")),
            H(1, "일정"),
            T(("단계", "기간", "우리가 할 일", "고객이 할 일")),
            H(1, "견적"),
            T(("항목", "수량", "금액"), rows=3),
            P("포함 범위와 제외 항목을 아래에 적습니다."),
            H(1, "요청"),
            P("이 문서를 읽고 무엇을 결정해 달라는 것인지 한 줄."),
        ),
    ),
)

def _east_asian(style, font: str) -> None:
    """The Korean face.

    `font.name` sets the Latin face only; a Hangul run reads `w:eastAsia` and
    falls back to Word's default without it — which is how a template that
    looks right in the styles pane comes out in the wrong face on the page.
    """
    style.font.name = font
    element = style.element.rPr.rFonts
    element.set(qn("w:eastAsia"), font)
    element.set(qn("w:cs"), font)


def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    if placeholder:
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = placeholder
        run.append(text)
        field.append(run)
    paragraph._p.append(field)


def build(spec: Spec) -> pathlib.Path:
    document = Document()
    accent = RGBColor.from_string(spec.accent)

    top, right, bottom, left = spec.margins_mm
    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(top)
        section.right_margin = Mm(right)
        section.bottom_margin = Mm(bottom)
        section.left_margin = Mm(left)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _field(footer, "PAGE", "1")

    normal = document.styles["Normal"]
    _east_asian(normal, spec.body_font)
    normal.font.size = Pt(spec.body_pt)
    normal.paragraph_format.space_after = Pt(6)
    # 160% — Korean needs more leading than Word's default gives it, and a
    # report set at 100% reads as a wall.
    normal.paragraph_format.line_spacing = 1.6

    for name, size, colour in (
        ("Title", spec.title_pt, accent),
        ("Heading 1", spec.h1_pt, accent),
        ("Heading 2", spec.h2_pt, None),
    ):
        style = document.styles[name]
        _east_asian(style, spec.heading_font)
        style.font.size = Pt(size)
        style.font.bold = True
        if colour is not None:
            style.font.color.rgb = colour
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    # 본문. Word calls it 본문 in Korean and it is where the writing goes —
    # `Normal` is the base everything inherits from, so changing it changes
    # headings and captions too. A document whose body cannot be restyled
    # without moving its headings is a document nobody restyles.
    body = document.styles["Body Text"]
    _east_asian(body, spec.body_font)
    body.font.size = Pt(spec.body_pt)
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.6

    subtitle = document.styles["Subtitle"]
    _east_asian(subtitle, spec.heading_font)
    subtitle.font.size = Pt(spec.body_pt + 1)
    subtitle.font.bold = False
    subtitle.font.color.rgb = RGBColor.from_string("666666")

    for name in ("List Bullet", "List Number", "Quote", "Caption"):
        if name in [s.name for s in document.styles]:
            _east_asian(document.styles[name], spec.body_font)
            document.styles[name].font.size = Pt(spec.body_pt)

    # ── the form's own two styles ──────────────────────────────────────────
    #
    # A blank form is mostly guidance — a line under each heading saying what
    # belongs there — and guidance has to be one thing the reader can select,
    # restyle or delete in a single move. Written as direct formatting it is
    # none of those: it is grey italic text that looks like guidance and
    # behaves like body, and typing over it leaves the new sentence grey.
    #
    # `next_paragraph_style` is the other half. Pressing Enter at the end of a
    # guidance line lands in 본문, so the form stops being grey the moment
    # somebody starts writing in it.
    guide = document.styles.add_style("안내", WD_STYLE_TYPE.PARAGRAPH)
    guide.base_style = document.styles["Body Text"]
    _east_asian(guide, spec.body_font)
    guide.font.size = Pt(spec.body_pt - 0.5)
    guide.font.italic = True
    guide.font.color.rgb = RGBColor.from_string("7A7A7A")
    guide.paragraph_format.space_after = Pt(2)
    guide.next_paragraph_style = document.styles["Body Text"]

    guide_list = document.styles.add_style("안내 목록", WD_STYLE_TYPE.PARAGRAPH)
    guide_list.base_style = document.styles["List Bullet"]
    _east_asian(guide_list, spec.body_font)
    guide_list.font.size = Pt(spec.body_pt - 0.5)
    guide_list.font.italic = True
    guide_list.font.color.rgb = RGBColor.from_string("7A7A7A")
    guide_list.paragraph_format.space_after = Pt(2)
    guide_list.next_paragraph_style = document.styles["Body Text"]

    # Table headers, for the same reason: a column name is a role, and a role
    # written as `bold=True` on every cell is twelve places to change one
    # decision.
    head = document.styles.add_style("표 머리", WD_STYLE_TYPE.PARAGRAPH)
    head.base_style = document.styles["Body Text"]
    _east_asian(head, spec.heading_font)
    head.font.size = Pt(spec.body_pt - 0.5)
    head.font.bold = True
    head.paragraph_format.space_after = Pt(0)
    head.paragraph_format.line_spacing = 1.2

    cell = document.styles.add_style("표 본문", WD_STYLE_TYPE.PARAGRAPH)
    cell.base_style = document.styles["Body Text"]
    _east_asian(cell, spec.body_font)
    cell.font.size = Pt(spec.body_pt - 0.5)
    cell.paragraph_format.space_after = Pt(0)
    cell.paragraph_format.line_spacing = 1.2

    # The template ships empty. Its value is the styles and the page, and a
    # paragraph of sample text left in one would arrive at the top of every
    # document written from it.
    for paragraph in list(document.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)

    out = ROOT / spec.folder / "template.docx"
    document.save(out)
    return out


def build_form(spec: Spec) -> pathlib.Path | None:
    """The blank form somebody downloads and fills in by hand.

    Built on the same document `build` just produced, so the styles, the page
    and the theme are the 서식's own — a form that looked like the export was
    the point. Kept as a second file because the export writer *appends* to
    `template.docx`, and a heading left in that one arrives at the top of every
    report written from it.
    """
    if not spec.form:
        return None
    document = Document(str(ROOT / spec.folder / "template.docx"))

    # Every line here names a style and sets nothing itself. A form built out
    # of direct formatting is a form that cannot be restyled: the reader who
    # wants their organisation's face has to walk every paragraph, and the one
    # who wants the guidance gone has to find each grey line by eye.
    for block in spec.form:
        if block[0] == "h":
            _, level, text = block
            # 제목 / 제목 1 / 제목 2 — built-ins, so Word shows them under
            # those names and the navigation pane and 목차 both work.
            document.add_heading(text, level=level)
        elif block[0] == "p":
            document.add_paragraph(block[1], style="안내")
            # Somewhere to write that is not the guidance. Typing over a
            # guidance line leaves the sentence in the guidance style, so the
            # form gives 본문 its own empty paragraph underneath.
            document.add_paragraph("", style="Body Text")
        elif block[0] == "b":
            for item in block[1]:
                document.add_paragraph(item, style="안내 목록")
            document.add_paragraph("", style="Body Text")
        elif block[0] == "t":
            _, columns, rows = block
            # `Table Grid` so the empty cells are visible. A form whose ruling
            # only appears once there is text in it is a form nobody can see
            # the shape of, which is the one thing a blank form is for.
            table = document.add_table(rows=rows + 1, cols=len(columns))
            table.style = "Table Grid"
            for index, name in enumerate(columns):
                paragraph = table.rows[0].cells[index].paragraphs[0]
                paragraph.text = name
                paragraph.style = document.styles["표 머리"]
            for row in table.rows[1:]:
                for used in row.cells:
                    used.paragraphs[0].style = document.styles["표 본문"]
            document.add_paragraph("", style="Body Text")

    out = ROOT / spec.folder / "form.docx"
    document.save(out)
    return out


def main() -> int:
    for spec in _SPECS:
        folder = ROOT / spec.folder
        if not folder.is_dir():
            print(f"없는 서식: {spec.folder}", file=sys.stderr)
            return 1
        path = build(spec)
        form = build_form(spec)
        size = f"{path.stat().st_size:>7,}"
        blank = f" + 양식 {form.stat().st_size:,}바이트" if form else ""
        print(f"{spec.folder:14} {size}바이트{blank}  {spec.note}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
