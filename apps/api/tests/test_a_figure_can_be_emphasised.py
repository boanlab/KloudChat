"""KPI strips and procedures reach all three export files as text, not pictures."""

from __future__ import annotations

import io
import re
import zipfile

import pytest

from app.services import report_export

SECTION = {
    "id": "s1",
    "heading": "성과",
    "content": (
        "도입 뒤 지표는 아래와 같다.\n\n"
        "```kpi\n"
        "32% | 오탐 감소\n"
        "1.4초 | 평균 응답\n"
        "99.2% | 가용성\n"
        "```\n\n"
        "오탐이 준 것이 가장 크다.\n"
    ),
}
REPORT = ("분기 보고", [SECTION])


def test_the_fence_is_read_as_a_strip_and_not_as_prose() -> None:
    lines = report_export._markdown_to_lines(SECTION["content"])
    kpi = [line for line in lines if line[0] == "kpi"]
    assert len(kpi) == 1
    assert kpi[0][1] == [("32%", "오탐 감소"), ("1.4초", "평균 응답"), ("99.2%", "가용성")]
    # No backticks survive in the prose.
    assert not any("```" in str(line[1]) for line in lines)


def test_a_line_without_a_label_is_dropped() -> None:
    # A figure needs a label.
    assert report_export._kpi_rows("32%\n1.4초 | 평균 응답") == [("1.4초", "평균 응답")]


def test_five_figures_are_cut_to_four() -> None:
    # At most four figures per strip.
    rows = report_export._kpi_rows("\n".join(f"{n} | 이름{n}" for n in range(5)))
    assert len(rows) == 4


def test_word_gets_a_real_table_not_a_picture() -> None:
    from docx import Document

    data = report_export.to_docx(*REPORT)
    document = Document(io.BytesIO(data))
    assert document.tables, "강조 수치가 표로 그려지지 않았다"
    table = document.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["32%", "1.4초", "99.2%"]
    assert [c.text for c in table.rows[1].cells] == ["오탐 감소", "평균 응답", "가용성"]
    run = table.rows[0].cells[0].paragraphs[0].runs[0]
    assert run.bold and run.font.size.pt > 14
    # No picture: the figures stay editable text.
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        assert not [n for n in archive.namelist() if "/media/" in n]


def test_the_pdf_draws_it() -> None:
    data = report_export.to_pdf(*REPORT)
    assert data.startswith(b"%PDF")
    assert len(data) > 1000


def test_hangul_gets_a_real_table() -> None:
    data = report_export.to_hwpx(*REPORT)
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        section = archive.read("Contents/section0.xml").decode()
    assert "<hp:tbl" in section, "강조 수치가 hwpx 표로 그려지지 않았다"
    for text in ("32%", "오탐 감소", "가용성"):
        assert text in section


@pytest.mark.parametrize("source", ["", "   ", "|", " | ", "값만있음"])
def test_an_empty_fence_writes_nothing(source: str) -> None:
    # An empty fence leaves no empty box.
    lines = report_export._markdown_to_lines(f"앞\n\n```kpi\n{source}\n```\n\n뒤\n")
    assert not [line for line in lines if line[0] == "kpi"]


# ── 절차 ────────────────────────────────────────────────────────────────

STEPS = {
    "id": "s2",
    "heading": "진행 방법",
    "content": (
        "다음 순서로 진행한다.\n\n"
        "```steps\n"
        "자료 수집 | 공개 데이터와 내부 로그를 모은다\n"
        "정제 | 중복과 결측을 걸러낸다\n"
        "분석 | 세 가지 기준으로 견준다\n"
        "```\n"
    ),
}


def test_a_procedure_is_read_as_steps() -> None:
    lines = report_export._markdown_to_lines(STEPS["content"])
    steps = [line for line in lines if line[0] == "steps"]
    assert len(steps) == 1
    assert steps[0][1][0] == ("자료 수집", "공개 데이터와 내부 로그를 모은다")
    assert len(steps[0][1]) == 3


def test_a_procedure_may_run_to_eight_steps() -> None:
    # At most eight steps.
    rows = report_export._kpi_rows("\n".join(f"단계{n} | 설명" for n in range(12)), limit=8)
    assert len(rows) == 8


def test_word_numbers_the_steps_itself() -> None:
    from docx import Document

    document = Document(io.BytesIO(report_export.to_docx("방법", [STEPS])))
    assert document.tables
    table = document.tables[0]
    assert len(table.rows) == 3
    assert [r.cells[0].text for r in table.rows] == ["1", "2", "3"]
    assert "자료 수집" in table.rows[0].cells[1].text
    # Text, not a picture.
    with zipfile.ZipFile(io.BytesIO(report_export.to_docx("방법", [STEPS]))) as archive:
        assert not [n for n in archive.namelist() if "/media/" in n]


def test_the_other_two_formats_draw_the_steps() -> None:
    assert report_export.to_pdf("방법", [STEPS]).startswith(b"%PDF")
    with zipfile.ZipFile(io.BytesIO(report_export.to_hwpx("방법", [STEPS]))) as archive:
        section = archive.read("Contents/section0.xml").decode()
    assert "<hp:tbl" in section
    assert "자료 수집" in section


def test_an_unfilled_procedure_writes_nothing() -> None:
    lines = report_export._markdown_to_lines("앞\n\n```steps\n\n```\n\n뒤\n")
    assert not [line for line in lines if line[0] == "steps"]


# ── 손으로 고친 뒤에도 ────────────────────────────────────────────────────

_EDITED = (
    "<p>도입 뒤 지표는 아래와 같다.</p>"
    '<div class="kpi">'
    "<div><strong>32%</strong><span>오탐 감소</span></div>"
    "<div><strong>1.4초</strong><span>평균 응답</span></div>"
    "</div>"
    '<ol class="steps">'
    "<li><strong>자료 수집</strong> <span>공개 데이터를 모은다</span></li>"
    "<li><strong>정제</strong> <span>중복을 걸러낸다</span></li>"
    "</ol>"
)


def test_an_edited_section_keeps_both_blocks() -> None:
    """An edited (HTML) section round-trips strips and procedures through `richtext` as fences."""
    from app.services import richtext

    markdown = richtext.to_markdown(_EDITED)
    assert "```kpi" in markdown
    assert "32% | 오탐 감소" in markdown
    assert "```steps" in markdown
    assert "자료 수집 | 공개 데이터를 모은다" in markdown
    # Not renumbered into the text, and not flattened into a plain list.
    assert "1. 자료 수집" not in markdown

    # Back out as the blocks, not as prose.
    kinds = [line[0] for line in report_export._markdown_to_lines(markdown)]
    assert "kpi" in kinds and "steps" in kinds


def test_the_blocks_survive_the_sanitiser() -> None:
    from app.services import design_templates

    kept = design_templates.sanitise(_EDITED)
    assert 'class="kpi"' in kept
    assert 'class="steps"' in kept


def test_an_edit_does_not_delete_the_diagrams() -> None:
    """Editing a section keeps its mermaid and chart fences."""
    from app.services import richtext

    edited = (
        "<p>앞 문장.</p>"
        '<figure class="diagram" data-source="pie showData&#10;    &quot;장비&quot; : 28">'
        '<img src="data:image/png;base64,AAAA" alt="">'
        "</figure>"
        "<p>뒤 문장.</p>"
    )
    markdown = richtext.to_markdown(edited)
    assert "```mermaid" in markdown
    assert 'pie showData' in markdown
    # Quotes and newlines come back as themselves.
    assert '"장비" : 28' in markdown
    # The picture is not also written into the prose; exporters look it up by digest.
    assert "data:image/png" not in markdown

    kinds = [line[0] for line in report_export._markdown_to_lines(markdown)]
    assert kinds.count("diagram") == 1


def test_a_plain_pasted_picture_is_still_a_picture() -> None:
    # A `<figure>` with no source is not a diagram.
    from app.services import richtext

    markdown = richtext.to_markdown(
        '<figure class="diagram"><img src="data:image/png;base64,AAAA" alt=""></figure>'
    )
    assert markdown.startswith("![")


def test_hangul_columns_are_weighted_and_reach_the_margin() -> None:
    """`_hwpx_table` column widths are weighted by content and sum to the text width."""
    import re

    from app.services.report_export import _HWPX_TEXT_WIDTH

    markup = report_export._hwpx_table(
        [["", "단계", "내용"], ["1", "자료 수집", "공개 데이터를 모은다"]],
        widths=[1, 5, 12],
    )
    widths = [int(w) for w in re.findall(r'<hp:cellSz width="(\d+)"', markup)]
    first_row = widths[:3]
    assert sum(first_row) == _HWPX_TEXT_WIDTH
    assert first_row[0] < first_row[1] < first_row[2]
    # Every row uses the same widths.
    assert widths[3:6] == first_row


def test_the_figures_are_centred_in_hangul() -> None:
    markup = report_export._hwpx_table([["32%"], ["오탐 감소"]], cell_para_pr=6)
    assert 'paraPrIDRef="6"' in markup
    # Shape 6 is declared in the header.
    assert any(shape[0] == 6 for shape in report_export._HWPX_PARA_SHAPES)
    assert report_export._HWPX_PARA_SHAPES[6][1] == "CENTER"


def test_the_step_name_and_its_detail_stay_apart() -> None:
    # Not one sentence with a number in front.
    section = {
        "id": "s",
        "heading": "방법",
        "content": "```steps\n자료 수집 | 공개 데이터를 모은다\n```\n",
    }
    with zipfile.ZipFile(io.BytesIO(report_export.to_hwpx("방법", [section]))) as archive:
        body = archive.read("Contents/section0.xml").decode()
    assert "자료 수집 공개 데이터를 모은다" not in body
    assert "자료 수집" in body and "공개 데이터를 모은다" in body


def test_a_step_number_is_centred_and_its_sentence_is_not() -> None:
    """In the HWPX table the step number is centred and the sentence is not."""
    markup = report_export._hwpx_table(
        [["", "단계", "내용"], ["1", "자료 수집", "공개 데이터를 모은다"]],
        widths=[1, 5, 12],
        cell_para_pr=[6, 3, 3],
    )
    shapes = re.findall(r'<hp:p paraPrIDRef="(\d)"', markup)
    # `shapes[0]` is the enclosing paragraph; then head row and body row, three cells each.
    assert shapes[1:7] == ["6", "3", "3", "6", "3", "3"]


def test_the_procedure_written_to_a_file_centres_its_numbers() -> None:
    section = {
        "id": "s",
        "heading": "방법",
        "content": "```steps\n자료 수집 | 공개 데이터를 모은다\n정제 | 걸러낸다\n```\n",
    }
    with zipfile.ZipFile(io.BytesIO(report_export.to_hwpx("방법", [section]))) as archive:
        body = archive.read("Contents/section0.xml").decode()
    assert 'paraPrIDRef="6"' in body
