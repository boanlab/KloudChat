"""Report export to .docx, .pdf and .hwpx: tables, page size, TOC, pictures, diagrams, forms."""

from __future__ import annotations

import base64
import io
import pathlib
import re
import zipfile
from xml.etree import ElementTree

from pypdf import PdfReader

from app.services import design_templates, pictures, report_export, richtext

_TABLE = (
    "본문 한 문단입니다.\n\n"
    "| 기준 | 외부 API | 온프레미스 |\n"
    "| --- | --- | --- |\n"
    "| 초기 비용 | 0원 | 약 3억 5천만 원 |\n"
    "| 데이터 보안 | 중위험 | 완벽 격리 |\n\n"
    "- 목록도 하나\n"
)
SECTIONS = [{"heading": "현황", "content": _TABLE}, {"heading": "결론", "content": "마무리."}]


def _docx(sections=SECTIONS) -> zipfile.ZipFile:
    return zipfile.ZipFile(io.BytesIO(report_export.to_docx("교체 검토", sections)))


# ── the parser the three exporters share ───────────────────────────────


def test_a_gfm_table_is_read_as_a_table():
    kinds = [kind for kind, _, _, _d in report_export._markdown_to_lines(_TABLE)]
    assert kinds == ["body", "table", "bullet"]


def test_the_rule_under_the_head_is_not_a_row():
    grid = next(t for k, t, _, _d in report_export._markdown_to_lines(_TABLE) if k == "table")
    assert grid.flat()[0] == ["기준", "외부 API", "온프레미스"]
    assert len(grid.rows) == 3


def test_an_escaped_pipe_stays_inside_its_cell():
    grid = next(
        t
        for k, t, _, _d in report_export._markdown_to_lines("| a\\|b | c |\n| --- | --- |\n")
        if k == "table"
    )
    assert grid.flat()[0] == ["a|b", "c"]


# ── .docx ──────────────────────────────────────────────────────────────


def test_the_docx_carries_a_real_table():
    body = _docx().read("word/document.xml").decode()
    assert body.count("<w:tbl>") == 1
    assert "| 기준 |" not in body


def test_the_docx_is_a4():
    # A4 in twips: 11906 × 16838 (Letter is 12240 × 15840).
    body = _docx().read("word/document.xml").decode()
    assert "11906" in body and "16838" in body


def test_the_docx_numbers_its_pages():
    archive = _docx()
    footers = [name for name in archive.namelist() if name.startswith("word/footer")]
    assert footers, "꼬리말이 없습니다"
    assert any("PAGE" in archive.read(name).decode() for name in footers)


def test_page_settings_reach_docx_and_pdf():
    settings = {
        "header": "대외비 검토본",
        "footer": "전략기획실",
        "pageNumbers": "page-total",
        "firstPageHeader": True,
        "margins": {"top": 25, "right": 18, "bottom": 24, "left": 19},
    }
    archive = zipfile.ZipFile(
        io.BytesIO(report_export.to_docx("제목", SECTIONS, page_settings=settings))
    )
    headers = "".join(
        archive.read(name).decode()
        for name in archive.namelist()
        if name.startswith("word/header")
    )
    footers = "".join(
        archive.read(name).decode()
        for name in archive.namelist()
        if name.startswith("word/footer")
    )
    document = archive.read("word/document.xml").decode()
    assert "대외비 검토본" in headers
    assert "전략기획실" in footers and "PAGE" in footers and "NUMPAGES" in footers
    assert 'w:top="1417"' in document and 'w:left="1077"' in document

    pdf = report_export.to_pdf("제목", SECTIONS, page_settings=settings)
    text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(pdf)).pages)
    assert "대외비 검토본" in text and "전략기획실" in text


def test_reference_list_is_formatted_and_appended_for_export():
    sources = [
        {
            "ordinal": 1,
            "title": "검증된 자료",
            "author": "김연구",
            "publisher": "한국연구원",
            "year": "2026",
            "url": "https://example.org/paper",
        }
    ]
    sections = report_export.with_references(SECTIONS, sources, "APA")

    assert sections[-1]["heading"] == "참고문헌"
    assert "김연구. (2026). 검증된 자료" in sections[-1]["content"]
    assert "https://example.org/paper" in sections[-1]["content"]
    assert len(SECTIONS) == 2, "화면의 원본 절 목록을 변경하면 안 됩니다"


def test_incomplete_web_source_has_no_empty_citation_punctuation():
    source = {"ordinal": 3, "title": "제목만 있는 검색 결과", "url": "https://example.org"}

    assert report_export.citation_text(source, "APA") == (
        "제목만 있는 검색 결과. https://example.org"
    )
    assert report_export.citation_text(source, "IEEE").startswith(
        '[3] “제목만 있는 검색 결과.”'
    )


def test_the_docx_carries_a_table_of_contents_field_only_when_long():
    body = _docx().read("word/document.xml").decode()
    assert "TOC" not in body
    long_sections = [
        {"heading": f"{i}절", "content": "본문입니다. " * 160} for i in range(1, 8)
    ]
    docx = zipfile.ZipFile(io.BytesIO(report_export.to_docx("긴 문서", long_sections)))
    assert "TOC" in docx.read("word/document.xml").decode()
    # `updateFields` makes Word fill the TOC field on open.
    assert "updateFields" in docx.read("word/settings.xml").decode()


def test_a_short_docx_does_not_leave_a_mostly_empty_cover_page():
    """No forced page break after the TOC."""
    body = _docx().read("word/document.xml").decode()
    assert 'w:type="page"' not in body


def test_a_designed_docx_table_uses_the_documents_accent_rule():
    body = zipfile.ZipFile(
        io.BytesIO(report_export.to_docx("제목", SECTIONS, tokens={"accent": "#0f766e"}))
    ).read("word/document.xml").decode()
    assert 'w:fill="F2F2F2"' in body
    assert 'w:color="0F766E"' in body
    assert '<w:pBdr><w:bottom w:val="single" w:sz="10" w:space="6" w:color="0F766E"' in body


# ── .pdf and .hwpx read the same document ──────────────────────────────


def test_the_pdf_is_a_pdf_and_has_the_table_in_it():
    data = report_export.to_pdf("교체 검토", SECTIONS)
    assert data[:5] == b"%PDF-"
    assert len(data) > len(report_export.to_pdf("교체 검토", [SECTIONS[1]]))


def test_pdf_page_total_is_the_actual_final_count():
    long_sections = [
        {"heading": f"긴 절 {number}", "content": ("검증할 본문입니다. " * 180)}
        for number in range(1, 8)
    ]
    pages = PdfReader(
        io.BytesIO(
            report_export.to_pdf(
                "쪽 수 검증",
                long_sections,
                page_settings={"pageNumbers": "page-total"},
            )
        )
    ).pages
    assert len(pages) > 1
    for number, page in enumerate(pages, 1):
        assert f"{number} / {len(pages)}" in (page.extract_text() or "")


def test_a_manual_page_break_survives_html_and_all_three_document_formats():
    html = (
        '<p>첫 쪽의 마지막 문단</p>'
        '<div data-page-break="true" class="page-break"></div>'
        '<p>둘째 쪽의 첫 문단</p>'
    )
    clean = design_templates.sanitise(html, editable_styles=True)
    assert 'data-page-break="true"' in clean
    markdown = richtext.to_markdown(clean)
    assert "<!-- pagebreak -->" in markdown
    sections = [{"heading": "수동 나눔", "content": markdown}]

    docx = zipfile.ZipFile(io.BytesIO(report_export.to_docx("제목", sections)))
    document_xml = docx.read("word/document.xml").decode()
    assert 'w:type="page"' in document_xml

    pdf = PdfReader(io.BytesIO(report_export.to_pdf("제목", sections)))
    assert len(pdf.pages) == 2
    assert "첫 쪽의 마지막 문단" in (pdf.pages[0].extract_text() or "")
    assert "둘째 쪽의 첫 문단" in (pdf.pages[1].extract_text() or "")

    hwpx = zipfile.ZipFile(io.BytesIO(report_export.to_hwpx("제목", sections)))
    section_xml = hwpx.read("Contents/section0.xml").decode()
    assert 'pageBreak="1"' in section_xml


def _hwpx(sections=SECTIONS) -> zipfile.ZipFile:
    return zipfile.ZipFile(io.BytesIO(report_export.to_hwpx("교체 검토", sections)))


def test_the_hwpx_carries_a_real_table():
    """The .hwpx draws a table as `hp:tbl`, not as pipes or `·`-joined lines."""
    body = _hwpx().read("Contents/section0.xml").decode()
    assert body.count("<hp:tbl") == 1
    assert body.count("<hp:tc ") == 9  # 3 columns × 3 rows
    assert "| 기준 |" not in body
    assert "기준 · 외부 API" not in body


def test_hwpx_page_geometry_and_first_page_flags_follow_page_settings():
    archive = zipfile.ZipFile(
        io.BytesIO(
            report_export.to_hwpx(
                "쪽 설정",
                SECTIONS,
                page_settings={
                    "firstPageHeader": True,
                    "pageNumbers": "none",
                    "margins": {"top": 25, "right": 19, "bottom": 21, "left": 17},
                },
            )
        )
    )
    body = archive.read("Contents/section0.xml").decode()
    assert 'hideFirstHeader="0"' in body
    assert 'hideFirstPageNum="1"' in body
    assert 'top="7087"' in body
    assert 'right="5386"' in body
    assert 'bottom="5953"' in body
    assert 'left="4819"' in body


def test_hwpx_page_furniture_is_editable_and_uses_a_live_page_number():
    archive = zipfile.ZipFile(
        io.BytesIO(
            report_export.to_hwpx(
                "제출 보고서",
                SECTIONS,
                page_settings={
                    "header": "사업 검토 자료",
                    "footer": "기획조정실",
                    "pageNumbers": "page-total",
                },
            )
        )
    )
    body = archive.read("Contents/section0.xml").decode()
    assert '<hp:header id="10001" applyPageType="BOTH">' in body
    assert "사업 검토 자료" in body
    assert '<hp:footer id="10002" applyPageType="BOTH">' in body
    assert "기획조정실" in body
    assert '<hp:pageNum pos="BOTTOM_RIGHT" formatType="DIGIT"' in body


def test_every_border_the_cells_point_at_is_defined():
    """Every `borderFillIDRef` a cell uses is defined in the header."""
    archive = _hwpx()
    head = archive.read("Contents/header.xml").decode()
    body = archive.read("Contents/section0.xml").decode()
    defined = set(re.findall(r'<hh:borderFill id="(\d+)"', head))
    used = set(re.findall(r'borderFillIDRef="(\d+)"', body))
    assert used, "표가 테두리를 참조하지 않습니다"
    assert not used - defined


def test_the_header_lists_its_parts_in_the_schema_order():
    """`refList` children follow the OWPML schema order; Hancom will not open a file otherwise."""
    head = _hwpx().read("Contents/header.xml").decode()
    order = re.findall(
        r"<hh:(fontfaces|borderFills|charProperties|paraProperties|styles)", head
    )
    assert order == ["fontfaces", "borderFills", "charProperties", "paraProperties", "styles"]


def test_every_part_of_the_hwpx_is_well_formed_xml():
    archive = _hwpx()
    for name in archive.namelist():
        if name.endswith((".xml", ".hpf")):
            ElementTree.fromstring(archive.read(name))


def test_a_table_it_cannot_render_falls_back_to_lines_rather_than_breaking():
    """A table too wide to render falls back to `·`-joined lines."""
    cells = " | ".join(str(i) for i in range(20))
    rule = " | ".join(["---"] * 20)
    wide = f"| {cells} |\n| {rule} |\n"
    body = _hwpx([{"heading": "가", "content": wide}]).read("Contents/section0.xml").decode()
    assert "<hp:tbl" not in body
    assert "0 · 1 · 2" in body


def test_every_format_opens():
    """Each exporter produces a well-formed file."""
    assert zipfile.ZipFile(io.BytesIO(report_export.to_docx("제목", SECTIONS))).testzip() is None
    assert zipfile.ZipFile(io.BytesIO(report_export.to_hwpx("제목", SECTIONS))).testzip() is None
    assert report_export.to_pdf("제목", SECTIONS)[:5] == b"%PDF-"


# ── the gap a model leaves between table rows ──────────────────────────

_LOOSE = (
    "앞 문장.\n\n"
    "| 특성 | 설명 |\n\n"
    "| :--- | :--- |\n\n"
    "| 확장성 | 노드를 늘린다. |\n\n"
    "뒤 문장."
)


def test_blank_lines_between_rows_are_closed():
    """`tidy_tables` closes blank lines between a table's rows."""
    assert "| 특성 | 설명 |\n| :--- | :--- |\n| 확장성" in richtext.tidy_tables(_LOOSE)


def test_the_prose_around_the_table_keeps_its_blank_lines():
    tidied = richtext.tidy_tables(_LOOSE)
    assert tidied.startswith("앞 문장.\n\n")
    assert tidied.endswith("|\n\n뒤 문장.")


def test_two_tables_do_not_merge_into_one():
    kinds = [
        k
        for k, _, _, _d in report_export._markdown_to_lines(
            "| a |\n| --- |\n\n문장\n\n| b |\n| --- |\n"
        )
        if k == "table"
    ]
    assert len(kinds) == 2


def test_the_exporters_read_a_loose_table_as_a_table_anyway():
    """The exporters read a table with blank lines between rows as one table."""
    grid = next(t for k, t, _, _d in report_export._markdown_to_lines(_LOOSE) if k == "table")
    assert grid.flat()[0] == ["특성", "설명"]
    assert len(grid.rows) == 2


# ── the 서식's Word half ────────────────────────────────────────────────


def test_every_document_template_ships_a_word_template():
    """Every document template has a `template.docx`."""
    for row in design_templates.all_templates():
        if row.kind != "document":
            continue
        assert row.docx_template, f"{row.id} 에 template.docx 가 없습니다"


def test_the_export_comes_out_in_the_template_it_was_written_in():
    chosen = design_templates.get("doc-report")
    themed = zipfile.ZipFile(
        io.BytesIO(report_export.to_docx("제목", SECTIONS, template=chosen.docx_template))
    )
    styles = themed.read("word/styles.xml").decode()
    assert "함초롬" in styles
    assert "2B4C7E" in styles


def test_a_template_that_cannot_be_read_still_exports():
    """A missing template path still exports with defaults."""
    data = report_export.to_docx("제목", SECTIONS, template="/nowhere/template.docx")
    assert zipfile.ZipFile(io.BytesIO(data)).testzip() is None


def test_a_templated_export_still_carries_the_table_and_the_page_size():
    chosen = design_templates.get("doc-brief")
    body = (
        zipfile.ZipFile(
            io.BytesIO(report_export.to_docx("제목", SECTIONS, template=chosen.docx_template))
        )
        .read("word/document.xml")
        .decode()
    )
    assert body.count("<w:tbl>") == 1
    assert "11906" in body and "16838" in body


# ── 그림 ────────────────────────────────────────────────────────────────

#: A 1×1 PNG.
_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)
_SRC = pictures.encode("image/png", _PNG)
_WITH_FIGURE = f"앞 문단.\n\n![구성도]({_SRC})\n\n뒤 문단."


def test_a_picture_pasted_into_the_body_survives_into_markdown():
    """A `<figure>` in the body becomes a Markdown image with its caption."""
    figure = f'<figure><img src="{_SRC}" alt=""><figcaption>구성도</figcaption></figure>'
    markup = f"<p>앞</p>{figure}<p>뒤</p>"
    out = richtext.to_markdown(markup)
    assert f"![구성도]({_SRC})" in out
    assert out.index("앞") < out.index("![") < out.index("뒤")


def test_the_shared_parser_reads_it_as_a_picture_in_place():
    kinds = [k for k, _, _, _d in report_export._markdown_to_lines(_WITH_FIGURE)]
    assert kinds == ["body", "image", "body"]


def test_a_remote_address_is_dropped_rather_than_fetched():
    """Exporters never fetch remote images."""
    kinds = [
        k
        for k, _, _, _d in report_export._markdown_to_lines("![x](https://example.com/a.png)")
    ]
    assert "image" not in kinds


def test_the_picture_reaches_all_three_files():
    sections = [{"heading": "구성", "content": _WITH_FIGURE}]
    archive = zipfile.ZipFile(io.BytesIO(report_export.to_docx("제목", sections)))
    assert [n for n in archive.namelist() if n.startswith("word/media/")]

    assert report_export.to_pdf("제목", sections)[:5] == b"%PDF-"

    hwpx = zipfile.ZipFile(io.BytesIO(report_export.to_hwpx("제목", sections)))
    body = hwpx.read("Contents/section0.xml").decode()
    # An HWPX picture needs the BinData bytes, the content.hpf item, the manifest
    # entry and the `<hp:pic>` reference to agree, or Hancom will not open the file.
    parts = hwpx.namelist()
    stored = [n for n in parts if n.startswith("BinData/")]
    assert len(stored) == 1, parts
    name = stored[0].split("/")[1].split(".")[0]
    assert f'binaryItemIDRef="{name}"' in body
    assert f'href="BinData/{name}' in hwpx.read("Contents/content.hpf").decode()
    assert f'odf:full-path="{stored[0]}"' in hwpx.read("META-INF/manifest.xml").decode()
    assert "구성도" in body
    assert "[그림]" not in body


# ── mermaid 다이어그램 ──────────────────────────────────────────────────

_MERMAID_SRC = "graph TD\n    A[방화벽] --> B[물리 서버]\n    B --> C[데이터베이스]"
_WITH_DIAGRAM = f"앞 문단.\n\n```mermaid\n{_MERMAID_SRC}\n```\n\n뒤 문단."


def test_a_mermaid_fence_is_a_diagram_and_not_prose():
    """A mermaid fence is read as a `diagram` line, not body."""
    kinds = [k for k, _, _, _d in report_export._markdown_to_lines(_WITH_DIAGRAM)]
    assert kinds == ["body", "diagram", "body"]


def test_the_key_ignores_trailing_whitespace():
    """`diagram_key` is stable under trailing whitespace."""
    assert report_export.diagram_key(_MERMAID_SRC) == report_export.diagram_key(
        _MERMAID_SRC + "  \n"
    )


def test_a_drawn_diagram_reaches_the_file_as_a_picture():
    key = report_export.diagram_key(_MERMAID_SRC)
    section = {
        "heading": "구성",
        "content": _WITH_DIAGRAM,
        "diagrams": {key: pictures.encode("image/png", _PNG)},
    }
    archive = zipfile.ZipFile(io.BytesIO(report_export.to_docx("제목", [section])))
    assert [n for n in archive.namelist() if n.startswith("word/media/")]
    body = archive.read("word/document.xml").decode()
    assert "graph TD" not in body


def test_an_undrawn_diagram_says_so_rather_than_leaving_a_gap():
    """An undrawn diagram leaves a placeholder line, not a gap."""
    body = (
        zipfile.ZipFile(
            io.BytesIO(report_export.to_docx("제목", [{"heading": "가", "content": _WITH_DIAGRAM}]))
        )
        .read("word/document.xml")
        .decode()
    )
    assert "다이어그램" in body
    assert "graph TD" not in body


def test_the_source_never_leaks_into_any_format():
    sections = [{"heading": "가", "content": _WITH_DIAGRAM}]
    hwpx = (
        zipfile.ZipFile(io.BytesIO(report_export.to_hwpx("제목", sections)))
        .read("Contents/section0.xml")
        .decode()
    )
    assert "graph TD" not in hwpx
    assert "[다이어그램]" in hwpx


def test_an_unclosed_fence_is_prose_rather_than_a_lost_paragraph():
    """An unclosed fence is exported as prose."""
    kinds = [k for k, _, _, _d in report_export._markdown_to_lines("앞\n\n```mermaid\ngraph TD\n")]
    assert "diagram" not in kinds
    assert kinds.count("body") >= 2


def test_a_drawn_diagram_is_embedded_in_hangul_too():
    """A drawn diagram is embedded in the .hwpx."""
    import base64

    from app.services import pictures

    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmM"
        "IQAAAABJRU5ErkJggg=="
    )
    section = {
        "heading": "구조",
        "content": _WITH_DIAGRAM,
        "diagrams": {
            report_export.diagram_key(_MERMAID_SRC): pictures.encode("image/png", png)
        },
    }
    hwpx = zipfile.ZipFile(io.BytesIO(report_export.to_hwpx("제목", [section])))
    body = hwpx.read("Contents/section0.xml").decode()
    assert [n for n in hwpx.namelist() if n.startswith("BinData/")]
    assert "binaryItemIDRef=" in body
    assert "[다이어그램]" not in body


def test_an_undrawn_diagram_still_says_so():
    """An undrawn diagram leaves a placeholder line in the .hwpx."""
    hwpx = zipfile.ZipFile(
        io.BytesIO(report_export.to_hwpx("제목", [{"heading": "구조", "content": _WITH_DIAGRAM}]))
    )
    assert "[다이어그램]" in hwpx.read("Contents/section0.xml").decode()
    assert not [n for n in hwpx.namelist() if n.startswith("BinData/")]


def test_the_blank_form_is_made_of_styles() -> None:
    """Every paragraph in a blank form names a style and carries no direct formatting."""
    from docx import Document

    for row in design_templates.all_templates():
        if row.kind != "document":
            continue
        form = pathlib.Path(row.docx_template).with_name("form.docx")
        assert form.is_file(), f"{row.id} 에 form.docx 가 없습니다"
        document = Document(str(form))

        named = {style.name for style in document.styles}
        for wanted in ("Body Text", "안내", "안내 목록", "표 머리", "표 본문"):
            assert wanted in named, f"{row.id}: {wanted} 스타일이 없습니다"

        # Enter after a guidance line continues in 본문.
        assert document.styles["안내"].next_paragraph_style.name == "Body Text"

        def bare(paragraph) -> list[str]:
            """Direct run formatting not named by a style."""
            found = []
            for run in paragraph.runs:
                if run.font.color is not None and run.font.color.rgb is not None:
                    found.append("색")
                if run.italic:
                    found.append("기울임")
                if run.bold:
                    found.append("굵게")
            return found

        walls = list(document.paragraphs)
        for table in document.tables:
            for line in table.rows:
                for cell in line.cells:
                    walls.extend(cell.paragraphs)
        for paragraph in walls:
            assert not bare(paragraph), (
                f"{row.id}: '{paragraph.text[:24]}' 에 직접 서식 {bare(paragraph)}"
            )
            assert paragraph.style.name != "Normal" or not paragraph.text.strip(), (
                f"{row.id}: '{paragraph.text[:24]}' 이 Normal 입니다 — 본문은 Body Text 로"
            )


def test_the_exported_body_is_styled_too() -> None:
    """`to_docx` writes body paragraphs in `_BODY`, never in `Normal`."""
    import inspect

    source = inspect.getsource(report_export.to_docx)
    assert 'document.add_paragraph()' not in source, "본문 문단이 스타일 없이 만들어집니다"
    assert "_BODY" in source


def test_the_form_speaks_the_template_s_own_language() -> None:
    """A blank form's headings all appear in the template's own text."""
    from docx import Document

    for row in design_templates.all_templates():
        if row.kind != "document":
            continue
        document = Document(str(pathlib.Path(row.docx_template).with_name("form.docx")))
        headings = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.style.name.startswith("Heading") and paragraph.text.strip()
        ]
        assert headings, f"{row.id}: 양식에 제목이 없습니다"

        spoken = " ".join((*row.checks, row.instructions))
        stray = [head for head in headings if head not in spoken]
        assert not stray, (
            f"{row.id}: 양식이 서식에 없는 말을 씁니다 — {stray}. "
            "미리보기·확인 항목·지시 가운데 한 곳에는 나오는 말이어야 합니다."
        )
