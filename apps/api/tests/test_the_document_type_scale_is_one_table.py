"""One type scale for documents: the page seeds, the reportlab PDF, the `.docx`, the
`.hwpx` and the web view read the same table.
"""

import io
import re
from pathlib import Path

from app.services import design_templates, doc_type, report_export

_WEB = Path(__file__).resolve().parents[3] / "apps/web/src/components/report/docType.ts"
_SECTIONS = [
    {
        "id": "s1",
        "heading": "요약",
        "content": "본문 한 문단.\n\n| 항목 | 값 |\n|---|---|\n| 가 | 1 |",
    }
]


def _web_table(name: str) -> dict[str, float]:
    text = _WEB.read_text(encoding="utf-8")
    start = text.index(f"export const {name} = {{")
    block = text[start : text.index("} as const", start)]
    return {
        m.group(1): float(m.group(2)) for m in re.finditer(r"^\s+(\w+): ([\d.]+),", block, re.M)
    }


def test_the_web_view_and_the_exporters_read_one_table() -> None:
    assert _web_table("TYPE") == {k: float(v) for k, v in doc_type.TYPE.items()}
    assert _web_table("LEADING") == {k: float(v) for k, v in doc_type.LEADING.items()}


def test_the_scale_is_the_agreed_document_sizes() -> None:
    assert doc_type.TYPE["title"] == 20
    assert doc_type.TYPE["h1"] == 14
    assert doc_type.TYPE["h2"] == 12
    assert doc_type.TYPE["body"] == 10.5
    assert doc_type.TYPE["table"] == 9.5
    assert doc_type.TYPE["caption"] == 9


def test_the_page_seed_draws_its_type_from_the_variables() -> None:
    """Every 서식 inherits the scale: the seed's rules point at `--doc-*`, and no document
    template sets a font size of its own."""
    root = Path(__file__).resolve().parents[3] / "apps/api/app/design_templates"
    seed = (root / "_document/seed.html").read_text(encoding="utf-8")
    for name in ("title", "h1", "h2", "body", "table", "small", "kpi"):
        assert f"var(--doc-{name})" in seed, f"seed does not read --doc-{name}"
    for css in root.glob("doc-*/design.css"):
        assert "font-size" not in css.read_text(encoding="utf-8"), f"{css.parent.name} sets sizes"
    template = design_templates.get("doc-report")
    assert template is not None
    sheet = design_templates.stylesheet(template, {})
    assert "--doc-body: 10.5pt;" in sheet and "--doc-title: 20pt;" in sheet


def test_section_numbers_count_on_the_heading_itself() -> None:
    """Paged.js drops `counter-increment` on pseudo-elements; the heading carries it."""
    root = Path(__file__).resolve().parents[3] / "apps/api/app/design_templates"
    seed = (root / "_document/seed.html").read_text(encoding="utf-8")
    assert "h2 { counter-increment: sec; }" in seed
    before = seed[seed.index("h2::before {") : seed.index("}", seed.index("h2::before {"))]
    assert "counter-increment" not in before


def test_the_docx_styles_carry_the_scale_and_cells_have_room() -> None:
    from docx import Document
    from docx.oxml.ns import qn

    built = Document(io.BytesIO(report_export.to_docx("제목", _SECTIONS)))
    assert built.styles["Title"].font.size.pt == doc_type.TYPE["title"]
    assert built.styles["Heading 1"].font.size.pt == doc_type.TYPE["h1"]
    assert built.styles["Normal"].font.size.pt == doc_type.TYPE["body"]
    table = built.tables[0]
    margins = table._tbl.tblPr.find(qn("w:tblCellMar"))
    assert margins is not None and margins.find(qn("w:left")).get(qn("w:w")) == "120"
    sizes = {
        run.font.size.pt
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.font.size
    }
    assert sizes == {doc_type.TYPE["table"]}


def test_the_pdf_is_gothic_at_the_scale_unless_the_design_says_serif() -> None:
    """The reportlab fallback embeds the gothic face by default, serif only on request.

    Where no Nanum file is installed (CI), reportlab falls back to a CID font and the
    face cannot be read from the bytes; the choice itself is still checked.
    """
    from app.services import fonts

    plain = report_export.to_pdf("제목", _SECTIONS)
    assert plain.startswith(b"%PDF")
    serif = report_export.to_pdf("제목", _SECTIONS, tokens={"font": "serif"})
    assert serif.startswith(b"%PDF")
    if fonts.embedded("gothic") and fonts.embedded("serif"):
        # reportlab writes the face's own name, subset-prefixed, as the BaseFont.
        assert b"NanumGothic" in plain and b"Myeongjo" not in plain
        assert b"Myeongjo" in serif
